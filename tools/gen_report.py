# -*- coding: utf-8 -*-
"""
Sinh bao cao do an tot nghiep "Mat AI" ra file Word (.docx).
Chay: python tools/gen_report.py
Output: BaoCao_DoAn_MatAI.docx (thu muc goc project)
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"
BASE = 13

doc = Document()

# ---------- Base style ----------
def _set_base_style():
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(BASE)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def _margins():
    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(3.0)
        s.right_margin = Cm(2.0)

_set_base_style()
_margins()

# ---------- Helpers ----------
def _font(run, size=BASE, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

def para(text="", size=BASE, bold=False, italic=False, align="just",
         space_after=6, space_before=0, indent_first=0.0, color=None):
    p = doc.add_paragraph()
    a = {"just": WD_ALIGN_PARAGRAPH.JUSTIFY, "center": WD_ALIGN_PARAGRAPH.CENTER,
         "left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.alignment = a
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(indent_first)
    if text:
        r = p.add_run(text)
        _font(r, size, bold, italic, color)
    return p

def h1(num, text):
    doc.add_page_break()
    p = doc.add_heading(level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    label = f"CHƯƠNG {num}. {text.upper()}" if num != "" else text.upper()
    r = p.add_run(label)
    _font(r, 16, bold=True, color=(0x1F, 0x38, 0x64))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    return p

def h2(text):
    p = doc.add_heading(level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    _font(r, 14, bold=True, color=(0x2E, 0x54, 0x96))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p

def h3(text):
    p = doc.add_heading(level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    _font(r, 13, bold=True, italic=True, color=(0x40, 0x40, 0x40))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        r = p.add_run(bold_prefix)
        _font(r, BASE, bold=True)
        r2 = p.add_run(text)
        _font(r2)
    else:
        r = p.add_run(text)
        _font(r)
    return p

def numlist(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    _font(r)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    # light gray shading
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F2F2")
    p.paragraph_format.element.get_or_add_pPr().append(shd)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10.5)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    return p

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    _font(r, 11.5, italic=True, color=(0x40, 0x40, 0x40))
    return p

def table(headers, rows, widths=None, caption_text=None, header_above=True):
    if caption_text and header_above:
        caption(caption_text)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        rp = hdr[i].paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = rp.add_run(htext)
        _font(r, 12, bold=True, color=(0xFF, 0xFF, 0xFF))
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            cp = cells[i].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cp.paragraph_format.space_after = Pt(2)
            r = cp.add_run(str(val))
            _font(r, 11.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    if caption_text and not header_above:
        caption(caption_text)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ============================================================
#  TRANG BÌA
# ============================================================
def cover():
    def c(text, size, bold=False, sp_after=6, sp_before=0, italic=False, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(sp_after)
        p.paragraph_format.space_before = Pt(sp_before)
        r = p.add_run(text)
        _font(r, size, bold, italic, color)
        return p

    c("[TÊN TRƯỜNG ĐẠI HỌC]", 14, bold=True, sp_before=6)
    c("[TÊN KHOA / VIỆN]", 13, bold=True, sp_after=40)
    c("ĐỒ ÁN TỐT NGHIỆP", 22, bold=True, sp_after=10, color=(0x1F, 0x38, 0x64))
    c("─────────", 14, sp_after=30)
    c("XÂY DỰNG ỨNG DỤNG ANDROID “MẮT AI”", 18, bold=True, sp_after=6,
      color=(0x1F, 0x38, 0x64))
    c("HỖ TRỢ NGƯỜI KHIẾM THỊ NHẬN DIỆN TIỀN VIỆT NAM ĐỒNG",
      15, bold=True, sp_after=6)
    c("VÀ ĐỌC THỰC ĐƠN BẰNG TRÍ TUỆ NHÂN TẠO", 15, bold=True, sp_after=44)

    # info block
    rows = [
        ("Sinh viên thực hiện:", "[Họ và tên sinh viên]"),
        ("Mã số sinh viên:", "[MSSV]"),
        ("Lớp / Khóa:", "[Lớp – Khóa]"),
        ("Ngành:", "Công nghệ thông tin"),
        ("Giảng viên hướng dẫn:", "[Học hàm, học vị – Họ tên GVHD]"),
    ]
    for label, val in rows:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(label + "  ")
        _font(r1, 13, bold=True)
        r2 = p.add_run(val)
        _font(r2, 13)
    c("", 12, sp_after=40)
    c("Hà Nội, năm 2026", 13, bold=True, italic=True, sp_before=30)

cover()

# ============================================================
#  LỜI CẢM ƠN
# ============================================================
doc.add_page_break()
para("LỜI CẢM ƠN", size=15, bold=True, align="center", space_after=14,
     color=(0x1F, 0x38, 0x64))
para("Trước hết, em xin gửi lời cảm ơn chân thành và sâu sắc tới [Học hàm, học vị – "
     "Họ tên GVHD] – giảng viên đã trực tiếp hướng dẫn, định hướng và tận tình góp ý "
     "trong suốt quá trình em thực hiện đồ án tốt nghiệp này. Những nhận xét quý báu "
     "của thầy/cô đã giúp em hoàn thiện cả về mặt kỹ thuật lẫn cách trình bày.",
     indent_first=1.0)
para("Em cũng xin cảm ơn quý thầy cô trong [Khoa/Viện] – [Tên trường] đã trang bị cho "
     "em nền tảng kiến thức vững chắc trong những năm học vừa qua, làm tiền đề để em "
     "có thể triển khai đề tài.", indent_first=1.0)
para("Đặc biệt, em xin tri ân cộng đồng người khiếm thị và Hội Người mù – những người "
     "đã truyền cảm hứng để đề tài này ra đời, hướng tới một sản phẩm công nghệ thực sự "
     "mang lại giá trị cho cuộc sống.", indent_first=1.0)
para("Do thời gian và kiến thức còn hạn chế, đồ án không tránh khỏi những thiếu sót. "
     "Em rất mong nhận được sự góp ý của quý thầy cô để đề tài được hoàn thiện hơn.",
     indent_first=1.0)
para("Em xin chân thành cảm ơn!", italic=True, space_before=10)
para("Sinh viên thực hiện", align="right", bold=True, space_before=14, space_after=2)
para("[Họ và tên sinh viên]", align="right", italic=True)

# ============================================================
#  TÓM TẮT
# ============================================================
doc.add_page_break()
para("TÓM TẮT ĐỒ ÁN", size=15, bold=True, align="center", space_after=14,
     color=(0x1F, 0x38, 0x64))
para("Người khiếm thị tại Việt Nam gặp nhiều khó khăn trong các hoạt động thường ngày, "
     "đặc biệt là nhận diện mệnh giá tiền mặt và đọc thực đơn tại nhà hàng. Đồ án này "
     "xây dựng “Mắt AI” – một ứng dụng Android gốc (native Kotlin) ứng dụng Trí tuệ nhân "
     "tạo nhằm giải quyết hai bài toán đó, với triết lý thiết kế “dùng được khi nhắm mắt”.")
para("Về mặt kỹ thuật, ứng dụng tích hợp hai mô-đun lõi: (1) Nhận diện tiền Việt Nam đồng "
     "theo thời gian thực bằng mô hình phát hiện đối tượng YOLOv10n chạy hoàn toàn trên "
     "thiết bị (on-device) qua TensorFlow Lite, hỗ trợ đếm cộng dồn nhiều tờ; (2) Đọc thực "
     "đơn bằng mô hình ngôn ngữ – thị giác (Vision-Language Model) Llama 4 Scout qua Groq "
     "Cloud, trích xuất trực tiếp danh sách “món – giá” từ ảnh và đọc to bằng giọng nói "
     "tiếng Việt tự nhiên. Toàn bộ trải nghiệm được thiết kế quanh cử chỉ chạm/vuốt và "
     "lệnh giọng nói, kết hợp phản hồi rung và TTS, tối ưu cho người không nhìn thấy màn hình.")
para("Kết quả, đồ án đã hoàn thiện một ứng dụng chạy được trên thiết bị thật (phiên bản "
     "0.9.0), kèm quy trình huấn luyện mô hình tái lập được trên Google Colab và bộ tài "
     "liệu kỹ thuật đầy đủ. Báo cáo trình bày cơ sở lý thuyết, phân tích – thiết kế, quá "
     "trình triển khai, kiểm thử và đánh giá, cùng định hướng phát triển trong tương lai.")
para("Từ khóa: hỗ trợ người khiếm thị, Android, YOLOv10, TensorFlow Lite, Vision-Language "
     "Model, Text-to-Speech, accessibility.", italic=True, space_before=8)

# ============================================================
#  MỤC LỤC (auto field)
# ============================================================
doc.add_page_break()
para("MỤC LỤC", size=15, bold=True, align="center", space_after=12,
     color=(0x1F, 0x38, 0x64))

def add_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar"); fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Nhấn Ctrl+A rồi F9 để cập nhật mục lục tự động."
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar); run._r.append(instrText); run._r.append(fldChar2)
    run._r.append(t); run._r.append(fldChar3)

add_toc()

# ============================================================
#  CHƯƠNG 1
# ============================================================
h1(1, "Tổng quan đề tài")

h2("1.1. Đặt vấn đề")
para("Theo thống kê của Tổ chức Y tế Thế giới (WHO) và Hội Người mù Việt Nam, nước ta "
     "hiện có khoảng hai triệu người khiếm thị ở các mức độ khác nhau. Mất hoặc suy giảm "
     "thị lực ảnh hưởng trực tiếp tới khả năng tự chủ của họ trong sinh hoạt hằng ngày, "
     "đặc biệt ở những tác vụ vốn dựa hoàn toàn vào thị giác.")
para("Hai tình huống tiêu biểu được khảo sát trong đề tài này là:")
bullet("người khiếm thị rất khó phân biệt mệnh giá các tờ tiền giấy/polymer Việt Nam đồng "
       "có kích thước và chất liệu gần giống nhau, dẫn tới rủi ro bị nhầm lẫn hoặc lừa gạt "
       "khi thanh toán, nhận tiền thừa;", bold_prefix="Nhận diện tiền: ")
bullet("thực đơn tại quán ăn, nhà hàng hầu như chỉ in chữ thường, không có chữ nổi Braille, "
       "khiến người khiếm thị phải phụ thuộc vào người đi cùng hoặc nhân viên để biết món "
       "và giá.", bold_prefix="Đọc thực đơn: ")
para("Sự phát triển mạnh mẽ của điện thoại thông minh, camera chất lượng cao và các mô "
     "hình Trí tuệ nhân tạo (AI) gọn nhẹ chạy được ngay trên thiết bị mở ra cơ hội xây "
     "dựng một “trợ lý thị giác” bỏ túi. Đó là động lực hình thành đề tài “Mắt AI”.")

h2("1.2. Mục tiêu đề tài")
para("Mục tiêu tổng quát: xây dựng một ứng dụng Android hỗ trợ người khiếm thị tự chủ "
     "hơn trong hai tác vụ nhận diện tiền và đọc thực đơn, với trải nghiệm có thể sử dụng "
     "hoàn toàn mà không cần nhìn màn hình.")
para("Các mục tiêu cụ thể:")
numlist("Nghiên cứu và huấn luyện mô hình AI nhận diện 9 mệnh giá tiền Việt Nam đồng đang "
        "lưu hành (1.000đ – 500.000đ), chạy theo thời gian thực trên thiết bị di động.")
numlist("Xây dựng mô-đun đọc thực đơn: từ ảnh chụp menu trích xuất danh sách món ăn kèm "
        "giá và đọc to bằng tiếng Việt.")
numlist("Thiết kế giao diện và mô hình tương tác “accessibility-first”: cử chỉ chạm/vuốt, "
        "lệnh giọng nói, phản hồi âm thanh và rung.")
numlist("Tích hợp chuyển văn bản thành giọng nói (TTS) tiếng Việt tự nhiên.")
numlist("Đóng gói ứng dụng thành tệp APK cài đặt được, kèm tài liệu hướng dẫn và quy trình "
        "huấn luyện mô hình tái lập được.")

h2("1.3. Đối tượng và phạm vi nghiên cứu")
h3("1.3.1. Đối tượng nghiên cứu")
bullet("Người khiếm thị sử dụng điện thoại Android làm người dùng cuối.")
bullet("Các kỹ thuật thị giác máy tính: phân loại ảnh, phát hiện đối tượng (object detection).")
bullet("Mô hình ngôn ngữ – thị giác (Vision-Language Model) cho bài toán hiểu ảnh tài liệu.")
bullet("Công nghệ TTS, nhận dạng giọng nói và chuẩn truy cập (accessibility) trên Android.")
h3("1.3.2. Phạm vi")
bullet("Nền tảng: Android gốc (Kotlin), tối thiểu Android 8.0 (API 26).")
bullet("Tiền tệ: 9 mệnh giá VND phổ biến; loại trừ mệnh giá 100đ, 200đ, 500đ rất hiếm dùng.")
bullet("Ngôn ngữ thực đơn: tiếng Việt (có dấu), menu in trên giấy.")
bullet("Đề tài tập trung vào hai chức năng lõi; các chức năng mô tả vật thể, lịch sử, cài "
       "đặt và điều khiển bằng giọng nói được phát triển ở mức hỗ trợ.")

h2("1.4. Phương pháp thực hiện")
para("Đề tài kết hợp nghiên cứu lý thuyết và phát triển thực nghiệm theo quy trình lặp:")
bullet("Khảo sát nhu cầu người dùng và các ứng dụng hiện có (Seeing AI, Lookout, Be My "
       "Eyes, Cash Reader).")
bullet("Thu thập/huấn luyện mô hình trên Google Colab; chuyển đổi sang định dạng TensorFlow "
       "Lite để nhúng vào ứng dụng.")
bullet("Lập trình ứng dụng theo kiến trúc MVVM, sử dụng Jetpack Compose, Hilt, Coroutines.")
bullet("Kiểm thử đơn vị (unit test) cho phần logic và kiểm thử trải nghiệm bằng TalkBack/"
       "bịt mắt mô phỏng người khiếm thị.")

h2("1.5. Ý nghĩa của đề tài")
para("Về mặt thực tiễn, “Mắt AI” hướng tới một sản phẩm miễn phí, dễ tiếp cận, góp phần "
     "nâng cao tính tự chủ và chất lượng sống cho người khiếm thị Việt Nam. Về mặt khoa "
     "học, đề tài là một trường hợp nghiên cứu cụ thể về việc kết hợp mô hình AI on-device "
     "(YOLOv10n) với mô hình đám mây (Vision-Language Model) trong cùng một ứng dụng di "
     "động, cũng như áp dụng các nguyên tắc thiết kế truy cập cho người không nhìn thấy.")

h2("1.6. Khảo sát các giải pháp hiện có")
table(
    ["Ứng dụng", "Hạn chế đối với người dùng Việt", "Hướng giải quyết của Mắt AI"],
    [
        ["Microsoft Seeing AI", "Không nhận diện được mệnh giá VND; giao diện và giọng đọc "
         "chủ yếu tiếng Anh", "Huấn luyện riêng mô hình cho VND, TTS tiếng Việt"],
        ["Google Lookout", "Đọc văn bản nhưng không phân tích cấu trúc “món – giá” của menu",
         "Dùng VLM trích xuất trực tiếp cặp món và giá"],
        ["Be My Eyes", "Phụ thuộc tình nguyện viên trực tuyến, cần mạng và có người rảnh",
         "Tự động hoàn toàn bằng AI"],
        ["Cash Reader", "Tập trung tiền tệ nước ngoài, hỗ trợ VND hạn chế",
         "Chuyên biệt cho tiền Việt Nam đồng"],
    ],
    widths=[4.0, 6.5, 5.0],
    caption_text="Bảng 1.1. So sánh Mắt AI với các giải pháp hiện có",
)

h2("1.7. Cấu trúc báo cáo")
para("Báo cáo gồm sáu chương: Chương 1 trình bày tổng quan đề tài; Chương 2 trình bày cơ "
     "sở lý thuyết; Chương 3 phân tích và thiết kế hệ thống; Chương 4 trình bày quá trình "
     "triển khai; Chương 5 kiểm thử và đánh giá; Chương 6 kết luận và hướng phát triển.")

# ============================================================
#  CHƯƠNG 2
# ============================================================
h1(2, "Cơ sở lý thuyết và công nghệ")

h2("2.1. Mạng nơ-ron tích chập (CNN)")
para("Mạng nơ-ron tích chập (Convolutional Neural Network – CNN) là kiến trúc nền tảng của "
     "thị giác máy tính hiện đại. CNN gồm các lớp tích chập (convolution) học các bộ lọc "
     "phát hiện đặc trưng cục bộ (cạnh, góc, kết cấu), các lớp gộp (pooling) giảm kích thước "
     "không gian, và các lớp kết nối đầy đủ ở cuối để phân loại. Nhờ cơ chế chia sẻ trọng "
     "số và tính bất biến tịnh tiến, CNN học được biểu diễn phân cấp của ảnh với số tham số "
     "nhỏ hơn nhiều so với mạng kết nối đầy đủ thuần túy.")
para("Trong đề tài, CNN là xương sống của cả mô hình phân loại tiền (MobileNetV3 ở giai đoạn "
     "đầu) lẫn mô hình phát hiện tiền (YOLOv10n ở giai đoạn hoàn thiện).")

h2("2.2. Bài toán phát hiện đối tượng và họ mô hình YOLO")
para("Khác với phân loại ảnh (chỉ trả về một nhãn cho toàn ảnh), phát hiện đối tượng "
     "(object detection) đồng thời xác định vị trí (khung bao – bounding box) và nhãn của "
     "từng đối tượng trong ảnh. Đây là bài toán phù hợp hơn cho việc đếm nhiều tờ tiền và "
     "định vị tờ tiền trong khung hình lộn xộn.")
para("YOLO (You Only Look Once) là họ mô hình phát hiện một giai đoạn (one-stage), dự đoán "
     "khung bao và xác suất lớp trong một lần lan truyền duy nhất, cho tốc độ cao phù hợp "
     "thời gian thực. Phiên bản YOLOv10 cải tiến quan trọng ở chỗ loại bỏ bước hậu xử lý "
     "NMS (Non-Maximum Suppression) nhờ cơ chế gán nhãn kép (dual label assignment), giúp "
     "giảm độ trễ suy luận. Biến thể YOLOv10n (nano) là phiên bản nhẹ nhất, chỉ vài triệu "
     "tham số, được thiết kế cho thiết bị tài nguyên hạn chế như điện thoại.")
table(
    ["Tiêu chí", "Phân loại ảnh (MobileNetV3)", "Phát hiện đối tượng (YOLOv10n)"],
    [
        ["Đầu ra", "Một nhãn / ảnh", "Nhiều khung bao + nhãn + độ tin cậy"],
        ["Đếm nhiều tờ", "Không hỗ trợ trực tiếp", "Hỗ trợ tự nhiên"],
        ["Tiền bị gấp/che", "Dễ nhầm khi nền phức tạp", "Định vị tốt hơn nhờ học vùng"],
        ["Vai trò trong đề tài", "Giải pháp ban đầu (v0.3)", "Giải pháp hoàn thiện (v0.4+)"],
    ],
    widths=[3.8, 5.8, 5.9],
    caption_text="Bảng 2.1. So sánh hai cách tiếp cận cho mô-đun nhận diện tiền",
)

h2("2.3. Học chuyển giao và lượng tử hóa mô hình")
para("Do bộ dữ liệu tiền Việt Nam tương đối nhỏ, đề tài áp dụng học chuyển giao (transfer "
     "learning): khởi tạo mô hình từ trọng số đã huấn luyện trên tập dữ liệu lớn (ImageNet "
     "cho MobileNetV3, COCO cho YOLOv10n) rồi tinh chỉnh (fine-tune) trên dữ liệu tiền. "
     "Cách này giúp hội tụ nhanh và tránh quá khớp (overfitting).")
para("Để mô hình chạy được trên điện thoại, mô hình PyTorch được xuất sang ONNX rồi chuyển "
     "đổi sang định dạng TensorFlow Lite (TFLite). TFLite hỗ trợ lượng tử hóa (quantization) "
     "– biểu diễn trọng số bằng số nguyên 8-bit (INT8) thay cho dấu phẩy động 32-bit (FP32) "
     "– giúp giảm kích thước tệp và tăng tốc suy luận, đánh đổi một phần độ chính xác. "
     "Phiên bản hiện tại dùng FP32 để giữ độ chính xác, có thể chuyển INT8 khi cần tối ưu.")

h2("2.4. Mô hình ngôn ngữ – thị giác (Vision-Language Model)")
para("Mô hình ngôn ngữ lớn (Large Language Model – LLM) là mạng nơ-ron dựa trên kiến trúc "
     "Transformer, huấn luyện trên khối lượng văn bản khổng lồ, có khả năng hiểu và sinh "
     "ngôn ngữ tự nhiên. Mô hình ngôn ngữ – thị giác (Vision-Language Model – VLM) mở rộng "
     "LLM bằng cách bổ sung bộ mã hóa hình ảnh, cho phép mô hình “nhìn” ảnh và trả lời/sinh "
     "văn bản dựa trên nội dung ảnh.")
para("Trong đề tài, mô-đun đọc thực đơn sử dụng VLM Llama 4 Scout (Meta) thông qua dịch vụ "
     "đám mây Groq. Thay vì OCR truyền thống rồi phân tích cú pháp (vốn làm “phẳng” bố cục "
     "nhiều cột và khó ghép đúng món với giá), VLM nhìn trực tiếp ảnh menu, hiểu quan hệ "
     "không gian giữa tên món và giá, và trả về kết quả có cấu trúc dạng JSON. Đây là điểm "
     "khác biệt then chốt so với kế hoạch ban đầu (xem Chương 4).")
para("Groq là nền tảng suy luận sử dụng phần cứng chuyên dụng (LPU) cho tốc độ phản hồi cao; "
     "gói miễn phí đủ cho mục đích trình diễn của đồ án.")

h2("2.5. Quang học ký tự (OCR)")
para("Nhận dạng ký tự quang học (Optical Character Recognition – OCR) chuyển hình ảnh chứa "
     "chữ thành văn bản máy đọc được. Google ML Kit Text Recognition v2 là thư viện OCR "
     "on-device, hỗ trợ tốt tiếng Việt có dấu. Trong giai đoạn đầu, đề tài dùng ML Kit để "
     "lấy văn bản thô từ menu rồi phân tích bằng biểu thức chính quy; tuy nhiên cách này bị "
     "hạn chế với menu nhiều cột, dẫn tới quyết định chuyển sang VLM (mục 2.4).")

h2("2.6. Chuyển văn bản thành giọng nói (TTS) và nhận dạng giọng nói")
para("Chuyển văn bản thành giọng nói (Text-to-Speech – TTS) là kênh đầu ra chính của ứng "
     "dụng với người khiếm thị. Đề tài sử dụng hai nguồn TTS: dịch vụ FPT.AI TTS (đám mây, "
     "giọng tiếng Việt tự nhiên) làm lựa chọn ưu tiên, và TextToSpeech tích hợp sẵn của "
     "Android làm phương án dự phòng khi không có mạng hoặc chưa cấu hình khóa. Chiều ngược "
     "lại, nhận dạng giọng nói (Speech-to-Text) dùng SpeechRecognizer của Android để hiện "
     "thực điều khiển bằng lệnh tiếng Việt.")

h2("2.7. Khả năng truy cập (Accessibility) trên Android")
para("Android cung cấp dịch vụ đọc màn hình TalkBack cùng các thuộc tính hỗ trợ truy cập "
     "(content description, focus order, vùng chạm tối thiểu). Các nguyên tắc thiết kế cho "
     "người khiếm thị áp dụng trong đề tài gồm: vùng chạm lớn (≥ 64dp), tương phản màu cao, "
     "phản hồi đa kênh (âm thanh + rung), và mô hình tương tác bằng cử chỉ đơn giản, nhất "
     "quán để người dùng ghi nhớ mà không cần nhìn.")

h2("2.8. Công nghệ phát triển ứng dụng")
table(
    ["Thành phần", "Công nghệ sử dụng"],
    [
        ["Ngôn ngữ & nền tảng", "Kotlin, Android (minSdk 26, target/compileSdk 36)"],
        ["Giao diện", "Jetpack Compose, Material 3, Navigation Compose"],
        ["Camera", "CameraX (Preview + ImageAnalysis)"],
        ["AI on-device", "TensorFlow Lite + TFLite Support"],
        ["OCR", "Google ML Kit Text Recognition v2"],
        ["AI đám mây", "Groq Cloud – Llama 4 Scout (Vision)"],
        ["TTS", "FPT.AI TTS + Android TextToSpeech"],
        ["Giọng nói", "Android SpeechRecognizer"],
        ["Tiêm phụ thuộc (DI)", "Hilt (Dagger)"],
        ["Bất đồng bộ", "Kotlin Coroutines + Flow"],
        ["Lưu trữ", "Room (lịch sử), DataStore (cài đặt)"],
        ["Huấn luyện mô hình", "Google Colab, PyTorch/Ultralytics → ONNX → TFLite"],
    ],
    widths=[5.0, 9.5],
    caption_text="Bảng 2.2. Tổng hợp công nghệ sử dụng trong đề tài",
)

# ============================================================
#  CHƯƠNG 3
# ============================================================
h1(3, "Phân tích và thiết kế hệ thống")

h2("3.1. Phân tích yêu cầu")
h3("3.1.1. Yêu cầu chức năng")
table(
    ["Mã", "Chức năng", "Mô tả"],
    [
        ["CN-01", "Nhận diện tiền", "Nhận diện mệnh giá tờ tiền theo thời gian thực, đọc "
         "to mệnh giá, cho phép đếm cộng dồn nhiều tờ và đọc tổng."],
        ["CN-02", "Đọc thực đơn", "Chụp ảnh menu, trích xuất danh sách món + giá, đọc lần "
         "lượt từng món, chọn món và tính tổng tiền."],
        ["CN-03", "Phản hồi giọng nói", "Mọi kết quả và hướng dẫn đều được đọc bằng TTS "
         "tiếng Việt."],
        ["CN-04", "Điều khiển bằng cử chỉ", "Chạm đôi, vuốt 4 hướng, giữ lâu, lắc máy để "
         "thao tác mà không cần nhìn."],
        ["CN-05", "Lệnh giọng nói", "Mở chức năng, đọc lại, dừng, đổi tốc độ bằng lời nói."],
        ["CN-06", "Lịch sử", "Lưu và phát lại 20 lần quét gần nhất."],
        ["CN-07", "Cài đặt", "Tùy chỉnh tốc độ đọc, giọng, tương phản, rung, tự động chụp."],
    ],
    widths=[1.6, 3.4, 9.5],
    caption_text="Bảng 3.1. Danh sách yêu cầu chức năng",
)
h3("3.1.2. Yêu cầu phi chức năng")
bullet("Khả năng truy cập: dùng được hoàn toàn khi không nhìn màn hình; tương thích TalkBack.")
bullet("Hiệu năng: nhận diện tiền theo thời gian thực (độ trễ suy luận mục tiêu < 200ms/khung).")
bullet("Độ tin cậy: có cơ chế dự phòng khi mô hình lỗi hoặc mất mạng.")
bullet("Tính khả chuyển: chạy trên Android 8.0 trở lên, không yêu cầu phần cứng đặc biệt.")
bullet("Bảo mật: khóa API lưu ở local.properties, không đưa vào mã nguồn công khai.")

h2("3.2. Tác nhân và ca sử dụng (Use Case)")
para("Hệ thống có một tác nhân chính là Người dùng khiếm thị. Các ca sử dụng tiêu biểu được "
     "tổ chức quanh hai chức năng lõi và các chức năng phụ trợ.")
table(
    ["Ca sử dụng", "Tác nhân", "Mô tả tóm tắt"],
    [
        ["Đếm tiền", "Người dùng", "Đưa tờ tiền vào camera → nghe mệnh giá → vuốt xuống "
         "chọn → nghe tổng."],
        ["Đọc menu", "Người dùng", "Hướng camera vào menu → tự chụp → nghe danh sách món "
         "→ điều hướng và chọn món."],
        ["Ra lệnh giọng nói", "Người dùng", "Giữ phím tăng âm → nói lệnh → hệ thống thực thi."],
        ["Xem lịch sử", "Người dùng", "Mở lịch sử → nghe lại kết quả cũ."],
        ["Cấu hình", "Người dùng", "Thay đổi tốc độ đọc, giọng, độ tương phản…"],
    ],
    widths=[3.5, 2.8, 8.2],
    caption_text="Bảng 3.2. Đặc tả tóm tắt các ca sử dụng",
)
caption("Hình 3.1. Sơ đồ ca sử dụng tổng quát (chèn sơ đồ Use Case tại đây)")

h2("3.3. Kiến trúc tổng thể")
para("Ứng dụng được tổ chức theo mẫu kiến trúc MVVM (Model – View – ViewModel) kết hợp tiêm "
     "phụ thuộc bằng Hilt. Tầng View (Jetpack Compose) chỉ hiển thị trạng thái và chuyển sự "
     "kiện; tầng ViewModel giữ logic nghiệp vụ và trạng thái dạng StateFlow; tầng dữ liệu "
     "gồm các mô-đun AI, TTS, lưu trữ. Mô hình AI nhận diện tiền chạy on-device, còn mô hình "
     "đọc menu gọi qua đám mây.")
code(
"            ┌───────────────────────────────────────────┐\n"
"            │           Tầng Giao diện (Compose)         │\n"
"            │   MoneyScreen · MenuScreen · Settings...   │\n"
"            └───────────────┬───────────────────────────┘\n"
"                            │ sự kiện / StateFlow\n"
"            ┌───────────────▼───────────────────────────┐\n"
"            │            Tầng ViewModel (MVVM)           │\n"
"            │   MoneyViewModel · MenuViewModel · ...     │\n"
"            └───┬───────────┬───────────────┬───────────┘\n"
"                │           │               │\n"
"      ┌─────────▼──┐  ┌─────▼──────┐  ┌─────▼──────────┐\n"
"      │ Mô-đun ML  │  │ TtsManager │  │  Lưu trữ        │\n"
"      │ YOLOv10n   │  │ FPT+Android│  │  Room/DataStore │\n"
"      │ (on-device)│  └────────────┘  └────────────────┘\n"
"      │ Groq VLM   │\n"
"      │ (cloud)    │\n"
"      └────────────┘\n"
)
caption("Hình 3.2. Sơ đồ kiến trúc tổng thể của ứng dụng Mắt AI")

h2("3.4. Thiết kế mô hình tương tác cho người khiếm thị")
para("Toàn bộ thao tác được quy về một bộ cử chỉ nhất quán giữa các màn hình, để người dùng "
     "ghi nhớ một lần và dùng ở mọi nơi:")
table(
    ["Cử chỉ", "Trong chế độ đếm tiền", "Trong chế độ đọc menu"],
    [
        ["Vuốt xuống", "Chọn (cộng) tờ tiền đang thấy", "Chọn món hiện tại"],
        ["Vuốt phải / trái", "—", "Món tiếp theo / món trước"],
        ["Vuốt lên", "Chuyển sang đọc menu", "Chuyển sang đếm tiền"],
        ["Chạm đôi", "Đọc tổng đã chọn", "Đọc lại danh sách đã chọn"],
        ["Giữ lâu", "Xóa hết, đếm lại", "Quét lại menu mới"],
        ["Lắc máy / Vol−", "Dừng đọc (panic stop)", "Dừng đọc (panic stop)"],
        ["Giữ Vol+", "Kích hoạt lệnh giọng nói", "Kích hoạt lệnh giọng nói"],
    ],
    widths=[3.2, 5.6, 5.7],
    caption_text="Bảng 3.3. Bộ cử chỉ điều khiển thống nhất",
)
para("Nguyên tắc thiết kế áp dụng: (1) phản hồi tức thì bằng âm thanh cho mọi thao tác; "
     "(2) hướng dẫn bằng giọng nói khi vào màn hình và sau kết quả đầu tiên; (3) nhắc nhở "
     "khi người dùng để yên quá lâu (idle reminder); (4) tránh đếm trùng bằng cơ chế “chờ "
     "nhấc tờ tiền ra” trước khi đếm tờ kế tiếp.")

h2("3.5. Thiết kế cơ sở dữ liệu cục bộ")
para("Lịch sử quét được lưu bằng Room với thực thể ScanHistoryEntity. Cài đặt người dùng "
     "lưu bằng Jetpack DataStore (Preferences).")
table(
    ["Trường", "Kiểu", "Ý nghĩa"],
    [
        ["id", "Long (PK)", "Khóa chính tự tăng"],
        ["type", "ScanType (MONEY/MENU)", "Loại lần quét"],
        ["content", "String", "Nội dung đã đọc bằng TTS"],
        ["timestamp", "Long", "Thời điểm quét (epoch millis)"],
    ],
    widths=[3.0, 4.5, 7.0],
    caption_text="Bảng 3.4. Cấu trúc bảng lịch sử quét (ScanHistoryEntity)",
)

# ============================================================
#  CHƯƠNG 4
# ============================================================
h1(4, "Triển khai hệ thống")

h2("4.1. Quá trình tiến hóa kiến trúc")
para("Một đặc điểm đáng chú ý của đề tài là kiến trúc được điều chỉnh qua nhiều phiên bản "
     "dựa trên kết quả thực nghiệm. Bảng dưới tóm tắt các mốc chính (theo lịch sử commit):")
table(
    ["Phiên bản", "Mô-đun tiền", "Mô-đun menu", "Lý do thay đổi"],
    [
        ["v0.3", "MobileNetV3 (phân loại, Kaggle)", "ML Kit OCR + parser regex",
         "Giải pháp khởi đầu, đơn giản"],
        ["v0.4", "YOLOv10n (phát hiện, Roboflow)", "Gemini 1.5 Flash (VLM)",
         "OCR mất bố cục cột; phân loại khó với tiền gấp/che"],
        ["v0.5–0.6", "YOLOv10n", "Groq Llama 4 Scout (Vision)",
         "VLM nhìn ảnh trực tiếp ghép “món–giá” chính xác hơn"],
        ["v0.7–0.9", "YOLOv10n + đếm cộng dồn", "Điều hướng + chọn món",
         "Hoàn thiện trải nghiệm cho người khiếm thị"],
    ],
    widths=[2.4, 4.2, 4.2, 4.0],
    caption_text="Bảng 4.1. Các mốc tiến hóa kiến trúc của Mắt AI",
)
para("Cấu trúc mã được thiết kế để cô lập sự thay đổi: mọi bộ nhận diện tiền đều cài đặt "
     "giao diện MoneyClassifier, cho phép thay mô hình mà không sửa tầng giao diện hay "
     "ViewModel.")

h2("4.2. Xây dựng và huấn luyện mô hình nhận diện tiền")
h3("4.2.1. Dữ liệu")
para("Bộ dữ liệu chính lấy từ Roboflow Universe, gồm ảnh tiền Việt Nam ở nhiều trạng thái "
     "thực tế: tờ phẳng, tờ gấp, tờ vò nhàu và tờ bị che một phần – đây là tình huống phổ "
     "biến khi người khiếm thị cầm tiền. Dữ liệu được gán nhãn theo định dạng phát hiện đối "
     "tượng (khung bao + lớp) cho 9 mệnh giá. Giai đoạn đầu, một bộ dữ liệu phân loại trên "
     "Kaggle (~2.250 ảnh sau khi loại 200đ/500đ) được dùng cho MobileNetV3.")
h3("4.2.2. Quy trình huấn luyện YOLOv10n")
para("Quy trình được đóng gói trong notebook Colab (03_money_yolov10n_train.ipynb):")
numlist("Tải dữ liệu từ Roboflow theo định dạng YOLO.")
numlist("Kiểm tra và ánh xạ lại thứ tự lớp đúng quy ước của ứng dụng "
        "(0 = 500.000đ, …, 8 = 1.000đ).")
numlist("Huấn luyện từ trọng số tiền huấn luyện yolov10n.pt: 100 epoch, kích thước ảnh "
        "640×640, batch tự động.")
numlist("Xuất mô hình tốt nhất sang TFLite (kèm NMS, FP32) và sinh tệp nhãn.")
numlist("Thả hai tệp vnd_yolov10n.tflite và vnd_yolov10n_labels.txt vào thư mục "
        "app/src/main/assets/ml/.")
table(
    ["Tham số", "Giá trị"],
    [
        ["Kiến trúc", "YOLOv10n (nano)"],
        ["Trọng số khởi tạo", "yolov10n.pt (tiền huấn luyện COCO)"],
        ["Số epoch", "100"],
        ["Kích thước ảnh đầu vào", "640 × 640"],
        ["Batch size", "Tự động (auto-batch)"],
        ["Số lớp", "9 mệnh giá VND"],
        ["Định dạng xuất", "TensorFlow Lite, FP32, tích hợp NMS"],
    ],
    widths=[6.0, 8.5],
    caption_text="Bảng 4.2. Cấu hình huấn luyện mô hình YOLOv10n",
)
h3("4.2.3. Suy luận trên thiết bị")
para("Lớp Yolov10MoneyDetector nạp mô hình bằng cơ chế ánh xạ bộ nhớ (memory-map) để tăng "
     "tốc, chuẩn hóa khung hình về 640×640 và giá trị điểm ảnh về [0,1]. Mô hình trả về "
     "danh sách phát hiện đã sắp theo độ tin cậy giảm dần; ứng dụng lấy phát hiện cao nhất, "
     "lọc theo ngưỡng tin cậy 0,70 và ánh xạ chỉ số lớp sang mệnh giá. Việc nhúng TFLite "
     "không nén (noCompress) cho phép nạp nhanh hơn khoảng ba lần.")
code(
"Khung hình CameraX (ImageAnalysis)\n"
"   → Yolov10MoneyDetector: resize 640×640, chuẩn hóa /255\n"
"   → Interpreter.run() → [1, MAX_DET, 6]  (x1,y1,x2,y2,conf,class)\n"
"   → lọc conf ≥ 0,70 → ánh xạ class → mệnh giá VND\n"
"   → MoneyViewModel: ổn định khung → đọc TTS mệnh giá\n"
)
caption("Hình 4.1. Luồng suy luận nhận diện tiền on-device")

h3("4.2.4. Cơ chế dự phòng nhiều tầng")
para("Mô-đun ML được thiết kế dự phòng để ứng dụng luôn chạy được kể cả khi thiếu mô hình:")
numlist("Yolov10MoneyDetector – mô hình chính (xử lý tiền gấp/vò/che).")
numlist("TfliteMoneyClassifier – mô hình MobileNetV3 (dự phòng cấp 1).")
numlist("FakeMoneyClassifier – bộ giả lập luân phiên 9 mệnh giá để kiểm thử giao diện khi "
        "chưa có mô hình.")

h2("4.3. Mô-đun đếm tiền và đọc số thành chữ")
para("MoneyViewModel hiện thực logic “đếm cộng dồn”: khi một tờ tiền được nhận diện ổn định "
     "(đủ số khung và độ tin cậy), ứng dụng đọc mệnh giá nhưng chưa cộng. Người dùng vuốt "
     "xuống để xác nhận cộng tờ đó vào tổng; sau đó hệ thống yêu cầu nhấc tờ tiền ra (chờ "
     "đủ số khung “trống”) trước khi cho phép đếm tờ kế tiếp, nhằm tránh đếm trùng.")
para("Để đọc số tiền tự nhiên, đề tài tự xây dựng hàm chuyển số sang chữ tiếng Việt "
     "(NumberToVietnamese) xử lý đầy đủ các trường hợp đặc biệt của tiếng Việt: “mốt” "
     "(21 → hai mươi mốt), “lăm” (25 → hai mươi lăm), “linh” (101 → một trăm linh một), "
     "nhóm số 0 ở giữa (1.000.001 → một triệu không trăm linh một). Hàm hỗ trợ tới hàng tỷ.")

h2("4.4. Mô-đun đọc thực đơn bằng VLM")
para("MenuViewModel điều phối luồng đọc menu. Sau khi chụp ảnh, ứng dụng gửi thẳng ảnh tới "
     "Groq Vision (Llama 4 Scout) – bỏ qua bước OCR vì OCR làm phẳng bố cục cột khiến việc "
     "ghép món với giá kém chính xác. Ảnh được thu nhỏ về tối đa 1024px, nén JPEG chất lượng "
     "85, mã hóa base64 và đính kèm cùng một prompt yêu cầu trả về JSON gồm danh sách "
     "{tên món, giá} đã chuẩn hóa.")
code(
'{"items":[{"name":"Phở Bò","price":"50000"},\n'
'          {"name":"Bún Chả","price":"45000"}]}\n'
)
caption("Hình 4.2. Định dạng JSON do VLM trả về cho mô-đun đọc menu")
para("Lớp GroqMenuAnalyzer xử lý gọi API có cơ chế thử lại với độ trễ tăng dần (exponential "
     "backoff) cho các lỗi tạm thời (HTTP 429/503), tách JSON khỏi văn bản thừa, và chuẩn "
     "hóa giá (ví dụ “50k”, “50.000”, “1.5tr” đều quy về số nguyên VND). Trải nghiệm đọc "
     "menu cho phép vuốt phải/trái để duyệt món, vuốt xuống để chọn món, và đọc tổng tiền "
     "các món đã chọn – đối xứng với mô-đun đếm tiền để người dùng dễ ghi nhớ.")

h2("4.5. Mô-đun TTS và điều khiển giọng nói")
para("RoutedTtsEngine đóng vai trò bộ định tuyến TTS thông minh: nếu người dùng chọn giọng "
     "FPT và có mạng + khóa API hợp lệ thì dùng FPT.AI cho giọng tự nhiên; ngược lại dùng "
     "Android TTS. Cách này đảm bảo chất lượng giọng đọc nhưng vẫn hoạt động ngoại tuyến. "
     "VoiceCommandService bao bọc SpeechRecognizer, nhận lệnh tiếng Việt và ánh xạ linh "
     "hoạt theo từ khóa (so khớp “chứa”) để bền với cách nói tự nhiên.")

h2("4.6. Tự động chụp dựa trên chất lượng khung hình")
para("FrameQualityAnalyzer phân tích nhanh từng khung hình CameraX để quyết định thời điểm "
     "tự chụp: ước lượng độ sáng (trung bình kênh Y) và độ nét (phương sai Laplacian trên "
     "mẫu con). Khi cả hai điều kiện đạt ngưỡng trong nhiều khung liên tiếp, hệ thống coi "
     "khung hình đã ổn định và kích hoạt chụp – giúp người khiếm thị không phải canh nút bấm.")

h2("4.7. Một số giao diện chính")
caption("Hình 4.3. Màn hình đếm tiền (chèn ảnh chụp màn hình)")
caption("Hình 4.4. Màn hình đọc menu (chèn ảnh chụp màn hình)")
caption("Hình 4.5. Màn hình cài đặt và lịch sử (chèn ảnh chụp màn hình)")
para("(Ghi chú: chèn ảnh chụp màn hình thực tế của ứng dụng vào các vị trí trên trước khi "
     "nộp báo cáo.)", italic=True)

# ============================================================
#  CHƯƠNG 5
# ============================================================
h1(5, "Kiểm thử và đánh giá")

h2("5.1. Môi trường và phương pháp kiểm thử")
para("Ứng dụng được kiểm thử trên thiết bị Android thật (Android 8.0 trở lên). Đề tài áp "
     "dụng ba mức kiểm thử: (1) kiểm thử đơn vị cho phần logic; (2) kiểm thử tích hợp/đầu "
     "cuối cho từng mô-đun; (3) kiểm thử trải nghiệm bằng TalkBack và mô phỏng bịt mắt để "
     "đánh giá theo góc nhìn người khiếm thị.")

h2("5.2. Kiểm thử đơn vị")
para("Phần logic chuyển số thành chữ tiếng Việt – vốn quan trọng vì sai một chữ sẽ khiến "
     "người dùng nghe nhầm số tiền – được phủ bằng bộ kiểm thử đơn vị (NumberToVietnameseTest) "
     "gồm nhiều nhóm trường hợp:")
table(
    ["Nhóm test", "Ví dụ kiểm tra", "Kết quả"],
    [
        ["9 mệnh giá VND", "200.000 → “hai trăm nghìn”", "Đạt"],
        ["Hàng đơn vị 0–9", "9 → “chín”", "Đạt"],
        ["10–19 dùng “mười”", "15 → “mười lăm”", "Đạt"],
        ["20–99 “mốt/lăm”", "21 → “hai mươi mốt”; 25 → “hai mươi lăm”", "Đạt"],
        ["100–999 “linh”", "101 → “một trăm linh một”", "Đạt"],
        ["Nhóm 0 ở giữa", "1.000.001 → “một triệu không trăm linh một”", "Đạt"],
        ["Hàng triệu, tỷ", "1.500.000.000 → “một tỷ năm trăm triệu”", "Đạt"],
        ["Tổng đếm thực tế", "650.000 → “Sáu trăm năm mươi nghìn đồng”", "Đạt"],
        ["Biên/ngoại lệ", "Số âm và > 999 tỷ ném ngoại lệ", "Đạt"],
    ],
    widths=[3.6, 7.4, 3.5],
    caption_text="Bảng 5.1. Kết quả kiểm thử đơn vị mô-đun đọc số thành chữ",
)
para("Ngoài ra còn có bộ kiểm thử cho hàm phân tích giá tiền trong menu "
     "(MenuOcrParserParseAmountTest).")

h2("5.3. Đánh giá mô-đun nhận diện tiền")
para("Mục tiêu đề ra là độ chính xác top-1 trên tập kiểm tra đạt trên 92% và độ trễ suy "
     "luận dưới 200ms/khung trên dòng chip tầm trung. Bảng dưới là khuôn đánh giá; các ô "
     "“[…]” cần điền giá trị đo thực tế sau khi chạy đánh giá cuối cùng trên thiết bị mục tiêu.")
table(
    ["Chỉ số", "Mục tiêu", "Kết quả đo (điền sau)"],
    [
        ["mAP@0.5 (phát hiện)", "≥ 0,90", "[…]"],
        ["Độ chính xác mệnh giá (top-1)", "> 92%", "[…]"],
        ["Độ trễ suy luận / khung", "< 200 ms", "[…]"],
        ["Nhận diện tiền gấp/che", "Hoạt động ổn định", "[…]"],
        ["Tỷ lệ báo nhầm (false positive)", "Thấp", "[…]"],
    ],
    widths=[5.6, 4.4, 4.5],
    caption_text="Bảng 5.2. Khuôn đánh giá mô-đun nhận diện tiền",
)
para("Khuyến nghị bổ sung ma trận nhầm lẫn (confusion matrix) giữa 9 mệnh giá và một vài "
     "ảnh minh họa kết quả phát hiện (đúng/sai) khi hoàn tất đo đạc.")

h2("5.4. Đánh giá mô-đun đọc thực đơn")
para("Mô-đun đọc menu được đánh giá theo độ chính xác trích xuất món và ghép giá. Đề tài đã "
     "chuẩn bị notebook đánh giá (02_menu_eval.ipynb) trên tập dữ liệu menu tiếng Việt "
     "(Viet-Menu, ~840 ảnh) với các chỉ số precision/recall/F1 cho việc phát hiện món và độ "
     "chính xác trích xuất giá.")
table(
    ["Chỉ số", "Giá trị (điền sau)"],
    [
        ["Số ảnh menu đánh giá", "840"],
        ["Precision phát hiện món", "[…]"],
        ["Recall phát hiện món", "[…]"],
        ["F1 phát hiện món", "[…]"],
        ["Độ chính xác trích xuất giá", "[…]"],
    ],
    widths=[7.5, 7.0],
    caption_text="Bảng 5.3. Khuôn đánh giá mô-đun đọc thực đơn",
)
para("Kèm theo nên phân tích 3–5 trường hợp lỗi tiêu biểu (ví dụ: món xuống dòng, menu "
     "viết tay, giá theo khoảng) và giải thích nguyên nhân.")

h2("5.5. Đánh giá trải nghiệm người dùng")
para("Trải nghiệm được kiểm tra bằng cách bật TalkBack và sử dụng ứng dụng trong điều kiện "
     "bịt mắt, đồng thời đề xuất mời 3–5 người khiếm thị thử nghiệm. Các tiêu chí gồm: thời "
     "gian từ lúc mở chức năng đến khi nghe kết quả (mục tiêu < 10 giây), khả năng hoàn thành "
     "tác vụ mà không cần trợ giúp, và điểm khả dụng theo thang SUS (System Usability Scale).")
table(
    ["Tiêu chí", "Mục tiêu", "Kết quả (điền sau)"],
    [
        ["Hoàn thành đếm tiền không cần nhìn", "100% người thử", "[…]"],
        ["Hoàn thành đọc menu không cần nhìn", "≥ 80% người thử", "[…]"],
        ["Thời gian tới kết quả", "< 10 giây", "[…]"],
        ["Điểm SUS", "≥ 70", "[…]"],
    ],
    widths=[6.5, 4.0, 4.0],
    caption_text="Bảng 5.4. Khuôn đánh giá trải nghiệm người dùng",
)

h2("5.6. Nhận xét chung")
para("Về mặt chức năng, ứng dụng đã chạy thông suốt hai luồng lõi trên thiết bị thật, với "
     "kiến trúc dự phòng nhiều tầng bảo đảm không “chết” khi thiếu mô hình hoặc mất mạng. "
     "Logic nhạy cảm (đọc số tiền) đã được kiểm thử đơn vị kỹ lưỡng. Các đánh giá định lượng "
     "về độ chính xác mô hình và trải nghiệm người dùng cần được hoàn tất ở bước cuối với "
     "dữ liệu đo thực tế để đưa vào báo cáo.")

# ============================================================
#  CHƯƠNG 6
# ============================================================
h1(6, "Kết luận và hướng phát triển")

h2("6.1. Kết quả đạt được")
para("Đề tài đã hoàn thành các mục tiêu đặt ra:")
bullet("Xây dựng hoàn chỉnh ứng dụng Android “Mắt AI” (phiên bản 0.9.0) chạy được trên "
       "thiết bị thật, đóng gói thành APK.")
bullet("Triển khai mô-đun nhận diện tiền VND theo thời gian thực bằng YOLOv10n on-device, "
       "hỗ trợ đếm cộng dồn nhiều tờ và đọc tổng.")
bullet("Triển khai mô-đun đọc thực đơn bằng mô hình ngôn ngữ – thị giác, trích xuất “món – "
       "giá” trực tiếp từ ảnh.")
bullet("Thiết kế mô hình tương tác accessibility-first nhất quán (cử chỉ, giọng nói, rung, "
       "TTS), dùng được khi không nhìn màn hình.")
bullet("Xây dựng quy trình huấn luyện mô hình tái lập được trên Colab và bộ tài liệu kỹ "
       "thuật đầy đủ; kiểm thử đơn vị cho logic trọng yếu.")

h2("6.2. Hạn chế")
bullet("Mô-đun đọc menu phụ thuộc kết nối mạng và dịch vụ đám mây bên thứ ba (Groq).")
bullet("Các số liệu đánh giá định lượng (độ chính xác, độ trễ, SUS) cần được đo và hoàn "
       "thiện với người khiếm thị thật.")
bullet("Bộ dữ liệu tiền còn hạn chế về đa dạng điều kiện thực tế; lớp “không phải tiền” "
       "(unknown) cần củng cố.")
bullet("Chưa hỗ trợ phát hiện tiền giả và mô tả vật thể xung quanh (mới ở mức định hướng).")

h2("6.3. Hướng phát triển")
numlist("Đưa mô-đun đọc menu chạy on-device bằng VLM gọn nhẹ để hoạt động ngoại tuyến.")
numlist("Mở rộng dữ liệu tiền (đa dạng ánh sáng, nền, độ mới cũ) và bổ sung phát hiện "
        "tiền giả qua đặc điểm bảo an.")
numlist("Bổ sung mô tả khung cảnh xung quanh (đối tượng, chướng ngại) hỗ trợ di chuyển.")
numlist("Hoàn thiện điều khiển bằng giọng nói rảnh tay và tích hợp trợ lý hội thoại.")
numlist("Tối ưu hiệu năng (lượng tử hóa INT8, GPU delegate) và phát hành trên Google Play.")
numlist("Tổ chức đánh giá người dùng quy mô lớn hơn cùng Hội Người mù để hoàn thiện UX.")

h2("6.4. Kết luận")
para("“Mắt AI” cho thấy việc kết hợp mô hình AI chạy trên thiết bị với mô hình đám mây, "
     "đặt trong một thiết kế tương tác lấy người khiếm thị làm trung tâm, có thể tạo ra một "
     "công cụ hữu ích và khả thi về kỹ thuật. Đề tài là nền tảng tốt để tiếp tục phát triển "
     "thành một sản phẩm hỗ trợ cộng đồng người khiếm thị Việt Nam.")

# ============================================================
#  TÀI LIỆU THAM KHẢO
# ============================================================
h1("", "Tài liệu tham khảo")
refs = [
    "A. Wang, et al., “YOLOv10: Real-Time End-to-End Object Detection,” arXiv:2405.14458, 2024.",
    "A. Howard, et al., “Searching for MobileNetV3,” Proc. IEEE/CVF ICCV, 2019.",
    "Meta AI, “The Llama 4 herd: native multimodal models,” 2025. [Trực tuyến]. "
    "Truy cập: https://ai.meta.com/blog/llama-4/",
    "Google, “TensorFlow Lite – On-device machine learning,” [Trực tuyến]. "
    "Truy cập: https://www.tensorflow.org/lite",
    "Google, “ML Kit Text Recognition v2,” [Trực tuyến]. "
    "Truy cập: https://developers.google.com/ml-kit/vision/text-recognition/v2",
    "Google, “Android Developers – Jetpack Compose, CameraX, Room, Hilt,” [Trực tuyến]. "
    "Truy cập: https://developer.android.com",
    "Groq, “GroqCloud Documentation,” [Trực tuyến]. Truy cập: https://console.groq.com/docs",
    "FPT.AI, “Text to Speech API,” [Trực tuyến]. Truy cập: https://fpt.ai/tts",
    "Ultralytics, “YOLO Documentation,” [Trực tuyến]. Truy cập: https://docs.ultralytics.com",
    "World Health Organization, “Blindness and vision impairment,” 2023. [Trực tuyến]. "
    "Truy cập: https://www.who.int",
    "J. Brooke, “SUS: A quick and dirty usability scale,” Usability Evaluation in Industry, 1996.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"[{i}] {r}")
    _font(run, 12)

# ---------- Save ----------
out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   "BaoCao_DoAn_MatAI.docx")
doc.save(out)
print("Saved:", out)
