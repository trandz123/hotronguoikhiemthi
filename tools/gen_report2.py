# -*- coding: utf-8 -*-
"""
Sinh bao cao do an "Mat AI" theo MUC LUC do nguoi dung cung cap (I..X).
Cac muc dac thu "nhan dien khuon mat" da duoc thay cho phu hop voi du an
(nhan dien tien VND + doc thuc don bang AI).
Output: BaoCao_DoAn_MatAI.docx (ghi de o thu muc goc).
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"
BASE = 13
doc = Document()

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DIAG = os.path.join(ROOT, "docs", "diagrams")
INDENT_CM = 1.27  # thut dau dong 1 tab (mac dinh Word = 1.27 cm)

# ---------------- base style ----------------
st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(BASE)
st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
pf = st.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.5
pf.space_after = Pt(6)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(3.0); s.right_margin = Cm(2.0)

# ---------------- helpers ----------------
def _font(run, size=BASE, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

def para(text="", size=BASE, bold=False, italic=False, align="just",
         space_after=6, space_before=0, indent_first=0.0, color=None):
    p = doc.add_paragraph()
    p.alignment = {"just": WD_ALIGN_PARAGRAPH.JUSTIFY, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(indent_first)
    elif align in ("just", "left"):
        p.paragraph_format.first_line_indent = Cm(INDENT_CM)
    if text:
        _font(p.add_run(text), size, bold, italic, color)
    return p

def h1(label):
    doc.add_page_break()
    p = doc.add_heading(level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _font(p.add_run(label.upper()), 16, bold=True, color=(0x1F, 0x38, 0x64))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)

def h2(label):
    p = doc.add_heading(level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _font(p.add_run(label), 14, bold=True, color=(0x2E, 0x54, 0x96))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)

def h3(label):
    p = doc.add_heading(level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _font(p.add_run(label), 13, bold=True, italic=True, color=(0x40, 0x40, 0x40))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        _font(p.add_run(bold_prefix), BASE, bold=True)
        _font(p.add_run(text))
    else:
        _font(p.add_run(text))

def numlist(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _font(p.add_run(text))

def code(text):
    # Cac khoi ASCII so do da duoc thay bang anh PNG -> bo qua khong render
    SKIP = ("Tầng Giao diện (Compose)", "Khung hình CameraX (ImageAnalysis",
            "Người dùng hướng camera vào menu", "→ MoneyScreen : đưa tờ tiền")
    if any(m in text for m in SKIP):
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F2")
    p.paragraph_format.element.get_or_add_pPr().append(shd)
    r = p.add_run(text)
    r.font.name = "Consolas"; r.font.size = Pt(10.5)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")

def caption(text):
    # Tu dong chen anh PNG truoc chu thich neu trung so hieu hinh
    FIG = {
        "Hình 3.1.": ("01_kientruc.png", 15.0),
        "Hình 3.2.": ("04_usecase.png", 13.0),
        "Hình 3.3.": ("02_luong_tien.png", 8.5),
        "Hình 3.4.": ("03_luong_menu.png", 9.0),
        "Hình 3.5.": ("05_tuantu.png", 16.0),
        "Hình 5.1.": ("06_trienkhai.png", 15.5),
    }
    for pre, (img, w) in FIG.items():
        if text.startswith(pre):
            path = os.path.join(DIAG, img)
            if os.path.exists(path):
                ip = doc.add_paragraph()
                ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ip.paragraph_format.space_before = Pt(6)
                ip.paragraph_format.space_after = Pt(2)
                ip.add_run().add_picture(path, width=Cm(w))
            break
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    _font(p.add_run(text), 11.5, italic=True, color=(0x40, 0x40, 0x40))

def table(headers, rows, widths=None, caption_text=None):
    if caption_text:
        caption(caption_text)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        rp = hdr[i].paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(rp.add_run(htext), 12, bold=True, color=(0xFF, 0xFF, 0xFF))
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            cp = cells[i].paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cp.paragraph_format.space_after = Pt(2)
            _font(cp.add_run(str(val)), 11.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ============================================================
#  TRANG BÌA
# ============================================================
def cover():
    def c(text, size, bold=False, sp_after=6, sp_before=0, italic=False, color=None):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(sp_after); p.paragraph_format.space_before = Pt(sp_before)
        _font(p.add_run(text), size, bold, italic, color)
    c("[TÊN TRƯỜNG ĐẠI HỌC]", 14, bold=True, sp_before=6)
    c("[TÊN KHOA / VIỆN]", 13, bold=True, sp_after=40)
    c("ĐỒ ÁN TỐT NGHIỆP", 22, bold=True, sp_after=10, color=(0x1F, 0x38, 0x64))
    c("─────────", 14, sp_after=30)
    c("XÂY DỰNG ỨNG DỤNG ANDROID “MẮT AI”", 18, bold=True, sp_after=6, color=(0x1F, 0x38, 0x64))
    c("HỖ TRỢ NGƯỜI KHIẾM THỊ NHẬN DIỆN TIỀN VIỆT NAM ĐỒNG", 15, bold=True, sp_after=6)
    c("VÀ ĐỌC THỰC ĐƠN BẰNG TRÍ TUỆ NHÂN TẠO", 15, bold=True, sp_after=44)
    for label, val in [
        ("Sinh viên thực hiện:", "[Họ và tên sinh viên]"),
        ("Mã số sinh viên:", "[MSSV]"),
        ("Lớp / Khóa:", "[Lớp – Khóa]"),
        ("Ngành:", "Công nghệ thông tin"),
        ("Giảng viên hướng dẫn:", "[Học hàm, học vị – Họ tên GVHD]"),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        _font(p.add_run(label + "  "), 13, bold=True); _font(p.add_run(val), 13)
    c("Hà Nội, năm 2026", 13, bold=True, italic=True, sp_before=44)
cover()

# ============================================================
#  MỤC LỤC
# ============================================================
doc.add_page_break()
para("MỤC LỤC", size=15, bold=True, align="center", space_after=12, color=(0x1F, 0x38, 0x64))
p = doc.add_paragraph(); run = p.add_run()
f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = 'TOC \\o "1-3" \\h \\z \\u'
f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
tt = OxmlElement("w:t"); tt.text = "Nhấn Ctrl+A rồi F9 để cập nhật mục lục tự động."
f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
for e in (f1, it, f2, tt, f3):
    run._r.append(e)

# ---- DANH MỤC TỪ VIẾT TẮT ----
doc.add_page_break()
para("DANH MỤC TỪ VIẾT TẮT", size=15, bold=True, align="center", space_after=12,
     color=(0x1F, 0x38, 0x64))
table(
    ["Từ viết tắt", "Tiếng Anh / Đầy đủ", "Diễn giải"],
    [
        ["AI", "Artificial Intelligence", "Trí tuệ nhân tạo"],
        ["ML", "Machine Learning", "Học máy"],
        ["DL", "Deep Learning", "Học sâu"],
        ["CNN", "Convolutional Neural Network", "Mạng nơ-ron tích chập"],
        ["YOLO", "You Only Look Once", "Họ mô hình phát hiện đối tượng một giai đoạn"],
        ["NMS", "Non-Maximum Suppression", "Khử khung bao trùng lặp"],
        ["VLM", "Vision-Language Model", "Mô hình ngôn ngữ – thị giác"],
        ["LLM", "Large Language Model", "Mô hình ngôn ngữ lớn"],
        ["OCR", "Optical Character Recognition", "Nhận dạng ký tự quang học"],
        ["TTS", "Text-to-Speech", "Chuyển văn bản thành giọng nói"],
        ["STT", "Speech-to-Text", "Nhận dạng giọng nói thành văn bản"],
        ["TFLite", "TensorFlow Lite", "Định dạng mô hình cho thiết bị di động"],
        ["ONNX", "Open Neural Network Exchange", "Định dạng trung gian trao đổi mô hình"],
        ["API", "Application Programming Interface", "Giao diện lập trình ứng dụng"],
        ["MVVM", "Model – View – ViewModel", "Mẫu kiến trúc phần mềm"],
        ["DI", "Dependency Injection", "Tiêm phụ thuộc"],
        ["UI / UX", "User Interface / Experience", "Giao diện / Trải nghiệm người dùng"],
        ["APK", "Android Package", "Tệp cài đặt ứng dụng Android"],
        ["VND", "Vietnam Dong", "Tiền Việt Nam đồng"],
        ["mAP", "mean Average Precision", "Độ chính xác trung bình (đánh giá phát hiện)"],
        ["IoU", "Intersection over Union", "Tỷ lệ giao trên hợp của khung bao"],
        ["SUS", "System Usability Scale", "Thang đo khả dụng hệ thống"],
        ["WCAG", "Web Content Accessibility Guidelines", "Bộ nguyên tắc khả năng truy cập"],
    ],
    widths=[3.0, 6.0, 5.5],
)

# ---- DANH MỤC BẢNG & HÌNH ----
doc.add_page_break()
para("DANH MỤC BẢNG", size=15, bold=True, align="center", space_after=8,
     color=(0x1F, 0x38, 0x64))
para("(Trong Word: References → Insert Table of Figures → chọn nhãn “Bảng” để sinh tự động "
     "danh mục bảng. Các bảng được đánh số theo chương: Bảng 2.1, Bảng 3.1, …)", italic=True)
para("DANH MỤC HÌNH ẢNH", size=15, bold=True, align="center", space_after=8, space_before=14,
     color=(0x1F, 0x38, 0x64))
para("(Tương tự: References → Insert Table of Figures → chọn nhãn “Hình”. Xem Phụ lục A để "
     "biết danh sách đầy đủ các hình cần chèn và vị trí.)", italic=True)

# ============================================================
#  I. MỞ ĐẦU
# ============================================================
h1("I. Mở đầu")

h2("1.1. Bối cảnh và lý do chọn đề tài")
para("Việt Nam hiện có khoảng hai triệu người khiếm thị ở nhiều mức độ. Khiếm khuyết thị "
     "giác ảnh hưởng trực tiếp đến khả năng tự chủ của họ trong sinh hoạt hằng ngày, đặc "
     "biệt ở những công việc vốn dựa hoàn toàn vào mắt. Hai tình huống nổi bật, gây bất tiện "
     "và rủi ro thường xuyên, là:")
bullet("các tờ tiền giấy và polymer Việt Nam đồng có kích thước, chất liệu khá giống nhau "
       "nên người khiếm thị rất khó phân biệt mệnh giá, dễ bị nhầm lẫn hoặc lừa gạt khi "
       "thanh toán, nhận tiền thừa;", bold_prefix="Nhận diện tiền: ")
bullet("thực đơn ở quán ăn, nhà hàng hầu như chỉ in chữ thường, không có chữ nổi Braille, "
       "buộc người khiếm thị phải nhờ người đi cùng hoặc nhân viên đọc giúp.",
       bold_prefix="Đọc thực đơn: ")
para("Cùng lúc đó, điện thoại thông minh ngày càng phổ biến, camera chất lượng cao và các "
     "mô hình Trí tuệ nhân tạo (AI) gọn nhẹ có thể chạy ngay trên thiết bị. Đây là cơ hội "
     "để xây dựng một “trợ lý thị giác” bỏ túi, biến chiếc điện thoại thành “đôi mắt” thay "
     "thế cho người khiếm thị. Xuất phát từ nhu cầu thực tế và tính khả thi về công nghệ, "
     "em chọn đề tài “Mắt AI – ứng dụng Android hỗ trợ người khiếm thị nhận diện tiền Việt "
     "Nam đồng và đọc thực đơn bằng Trí tuệ nhân tạo”.")
para("So với các ứng dụng nước ngoài hiện có (Microsoft Seeing AI, Google Lookout, Be My "
     "Eyes, Cash Reader), đa số chưa hỗ trợ tốt tiền VND và tiếng Việt, hoặc phụ thuộc tình "
     "nguyện viên trực tuyến. Đề tài hướng tới một giải pháp chuyên biệt cho người dùng Việt, "
     "tự động hoàn toàn bằng AI.")

h2("1.2. Mục tiêu đề tài")
para("Mục tiêu tổng quát: xây dựng ứng dụng Android giúp người khiếm thị tự chủ hơn trong "
     "việc nhận diện tiền và đọc thực đơn, với trải nghiệm dùng được hoàn toàn mà không cần "
     "nhìn màn hình. Các mục tiêu cụ thể gồm:")
numlist("Huấn luyện mô hình AI nhận diện 9 mệnh giá VND đang lưu hành (1.000đ – 500.000đ), "
        "chạy theo thời gian thực trên thiết bị di động.")
numlist("Xây dựng mô-đun đọc thực đơn: từ ảnh chụp menu trích xuất danh sách “món – giá” và "
        "đọc to bằng tiếng Việt.")
numlist("Thiết kế mô hình tương tác “accessibility-first”: cử chỉ chạm/vuốt, lệnh giọng nói, "
        "phản hồi âm thanh và rung.")
numlist("Tích hợp chuyển văn bản thành giọng nói (TTS) tiếng Việt tự nhiên.")
numlist("Đóng gói ứng dụng thành tệp APK cài đặt được kèm tài liệu và quy trình huấn luyện "
        "mô hình tái lập được.")

h2("1.3. Phạm vi và giới hạn")
h3("1.3.1. Phạm vi")
bullet("Nền tảng: Android gốc (Kotlin), tối thiểu Android 8.0 (API 26).")
bullet("Tiền tệ: 9 mệnh giá VND phổ biến (loại trừ 100đ, 200đ, 500đ rất hiếm dùng).")
bullet("Thực đơn: tiếng Việt có dấu, menu in trên giấy.")
bullet("Hai chức năng lõi (nhận diện tiền, đọc menu); các chức năng lịch sử, cài đặt, lệnh "
       "giọng nói ở mức hỗ trợ.")
h3("1.3.2. Giới hạn")
bullet("Mô-đun đọc menu cần kết nối mạng (gọi mô hình đám mây).")
bullet("Chưa hỗ trợ phát hiện tiền giả và mô tả vật thể xung quanh (mới ở mức định hướng).")
bullet("Số liệu đánh giá định lượng cần đo bổ sung với người khiếm thị thật.")

h2("1.4. Phương pháp nghiên cứu")
para("Đề tài kết hợp nghiên cứu lý thuyết và phát triển thực nghiệm theo quy trình lặp "
     "(iterative), bao gồm các bước:")
numlist("Khảo sát nhu cầu người dùng khiếm thị và phân tích các ứng dụng hỗ trợ hiện có để "
        "xác định khoảng trống cần giải quyết.")
numlist("Nghiên cứu cơ sở lý thuyết về thị giác máy tính, mô hình ngôn ngữ – thị giác, công "
        "nghệ TTS và các nguyên tắc thiết kế khả năng truy cập.")
numlist("Thu thập dữ liệu, huấn luyện và đánh giá mô hình AI trên Google Colab; chuyển đổi "
        "mô hình sang định dạng phù hợp với thiết bị di động.")
numlist("Lập trình ứng dụng theo kiến trúc MVVM; tích hợp mô hình AI, TTS và lưu trữ cục bộ.")
numlist("Kiểm thử đơn vị cho phần logic và kiểm thử trải nghiệm bằng TalkBack/bịt mắt mô "
        "phỏng người khiếm thị; điều chỉnh lặp lại dựa trên kết quả.")
para("Điểm đặc trưng của phương pháp là tính lặp: kiến trúc kỹ thuật của đề tài đã được "
     "điều chỉnh qua nhiều phiên bản (từ phân loại ảnh sang phát hiện đối tượng cho tiền; "
     "từ OCR + luật sang mô hình ngôn ngữ – thị giác cho menu) dựa trên kết quả thực nghiệm "
     "thực tế, thay vì cố định ngay từ đầu.")

h2("1.5. Khảo sát các giải pháp hiện có")
para("Trên thị trường đã có một số ứng dụng hỗ trợ người khiếm thị, tuy nhiên phần lớn chưa "
     "tối ưu cho người dùng Việt Nam, đặc biệt ở hai bài toán tiền VND và thực đơn tiếng "
     "Việt. Bảng 1.1 so sánh các giải pháp tiêu biểu với hướng tiếp cận của Mắt AI.")
table(
    ["Ứng dụng", "Hạn chế đối với người dùng Việt", "Hướng giải quyết của Mắt AI"],
    [
        ["Microsoft Seeing AI", "Không nhận diện mệnh giá VND; giao diện và giọng đọc chủ "
         "yếu tiếng Anh", "Huấn luyện riêng mô hình cho VND, TTS tiếng Việt"],
        ["Google Lookout", "Đọc văn bản nhưng không phân tích cấu trúc “món – giá” của menu",
         "Dùng VLM trích xuất trực tiếp cặp món và giá"],
        ["Be My Eyes", "Phụ thuộc tình nguyện viên trực tuyến, cần mạng và có người rảnh",
         "Tự động hoàn toàn bằng AI"],
        ["Cash Reader", "Tập trung tiền tệ nước ngoài, hỗ trợ VND hạn chế",
         "Chuyên biệt cho tiền Việt Nam đồng"],
    ],
    widths=[3.6, 6.4, 4.5],
    caption_text="Bảng 1.1. So sánh Mắt AI với các giải pháp hiện có",
)

h2("1.6. Bố cục báo cáo")
para("Báo cáo được tổ chức thành mười phần: Phần I trình bày mở đầu (bối cảnh, mục tiêu, "
     "phạm vi); Phần II trình bày cơ sở lý thuyết và công nghệ; Phần III phân tích hệ thống; "
     "Phần IV trình bày triển khai và thực nghiệm; Phần V phân tích các mô hình triển khai; "
     "Phần VI mô tả giao diện người dùng; Phần VII bàn về ứng dụng thực tế và khả năng mở "
     "rộng; Phần VIII nêu các khó khăn; Phần IX kết luận và hướng phát triển; Phần X liệt kê "
     "tài liệu tham khảo. Cuối báo cáo có phần Phụ lục.")

# ============================================================
#  II. CƠ SỞ LÝ THUYẾT
# ============================================================
h1("II. Cơ sở lý thuyết")

h2("2.1. Tổng quan về Trí tuệ nhân tạo (AI)")
para("Trí tuệ nhân tạo (Artificial Intelligence – AI) là lĩnh vực khoa học máy tính nghiên "
     "cứu cách tạo ra hệ thống có khả năng thực hiện những công việc vốn đòi hỏi trí tuệ con "
     "người: nhận thức hình ảnh, hiểu ngôn ngữ, suy luận, ra quyết định. Trong đề tài này, "
     "AI đóng vai trò “đôi mắt” và “giọng nói” của ứng dụng: nhìn và hiểu nội dung hình ảnh "
     "(tiền, thực đơn) rồi diễn đạt lại bằng lời nói cho người khiếm thị. Hai nhánh AI được "
     "sử dụng là thị giác máy tính (computer vision) và xử lý ngôn ngữ tự nhiên (qua mô hình "
     "ngôn ngữ – thị giác).")

para("Lịch sử AI trải qua nhiều giai đoạn: từ các hệ chuyên gia dựa trên luật (thập niên "
     "1970–1980), tới sự trỗi dậy của học máy thống kê (1990–2000), và bùng nổ học sâu từ "
     "khoảng năm 2012 khi mạng nơ-ron tích chập đạt kết quả vượt trội trong nhận dạng ảnh. "
     "Gần đây, các mô hình nền tảng (foundation models) quy mô lớn như mô hình ngôn ngữ lớn "
     "và mô hình đa phương thức (multimodal) tạo ra bước nhảy về khả năng hiểu ảnh và ngôn "
     "ngữ — chính là công nghệ mà mô-đun đọc menu của đề tài tận dụng.")
para("AI có nhiều nhánh: thị giác máy tính (computer vision), xử lý ngôn ngữ tự nhiên "
     "(NLP), nhận dạng tiếng nói, robot học, hệ khuyến nghị… Đề tài này tập trung chủ yếu "
     "vào thị giác máy tính (nhận diện tiền, đọc menu từ ảnh) và một phần xử lý tiếng nói "
     "(TTS, nhận dạng lệnh giọng nói).")

h2("2.2. Machine Learning và Deep Learning")
para("Học máy (Machine Learning – ML) là nhánh của AI cho phép máy tính “học” quy luật từ "
     "dữ liệu thay vì được lập trình tường minh từng quy tắc. Học sâu (Deep Learning) là "
     "tập con của ML, sử dụng mạng nơ-ron nhân tạo nhiều lớp để tự động học các đặc trưng "
     "phân cấp từ dữ liệu thô.")
para("Học máy thường được chia thành ba nhóm chính: học có giám sát (supervised learning – "
     "học từ dữ liệu đã gán nhãn, ví dụ ảnh tiền kèm mệnh giá), học không giám sát "
     "(unsupervised learning – tìm cấu trúc ẩn trong dữ liệu không nhãn), và học tăng cường "
     "(reinforcement learning – học qua thử–sai và phần thưởng). Cả hai mô hình nhận diện "
     "tiền và đọc menu trong đề tài đều thuộc nhóm học có giám sát.")
para("Mạng nơ-ron nhân tạo (Artificial Neural Network) mô phỏng cách các nơ-ron sinh học "
     "kết nối: mỗi nơ-ron nhận nhiều đầu vào, nhân với trọng số, cộng lại, đưa qua hàm kích "
     "hoạt phi tuyến (ReLU, Sigmoid…) rồi truyền đi. Khi xếp chồng nhiều lớp nơ-ron, mạng có "
     "thể xấp xỉ các hàm phức tạp. Học sâu chính là việc dùng các mạng nơ-ron “sâu” (nhiều "
     "lớp) để tự động học đặc trưng từ dữ liệu thô thay vì phải thiết kế đặc trưng thủ công.")
h3("2.2.1. Mạng nơ-ron tích chập (CNN)")
para("Mạng nơ-ron tích chập (Convolutional Neural Network – CNN) là kiến trúc nền tảng của "
     "thị giác máy tính hiện đại. CNN gồm các lớp tích chập học bộ lọc phát hiện đặc trưng "
     "cục bộ (cạnh, góc, kết cấu), lớp gộp (pooling) giảm kích thước, và lớp kết nối đầy đủ "
     "để phân loại. Nhờ chia sẻ trọng số và bất biến tịnh tiến, CNN học biểu diễn ảnh hiệu "
     "quả với ít tham số. CNN là xương sống của cả mô hình nhận diện tiền trong đề tài.")
h3("2.2.2. Học chuyển giao và lượng tử hóa mô hình")
para("Do bộ dữ liệu tiền VND tương đối nhỏ, đề tài áp dụng học chuyển giao (transfer "
     "learning): khởi tạo từ trọng số đã huấn luyện trên tập lớn (COCO/ImageNet) rồi tinh "
     "chỉnh trên dữ liệu tiền, giúp hội tụ nhanh và tránh quá khớp. Để chạy trên điện thoại, "
     "mô hình được chuyển sang định dạng TensorFlow Lite; kỹ thuật lượng tử hóa (quantization) "
     "có thể biểu diễn trọng số bằng số nguyên 8-bit (INT8) thay cho dấu phẩy động 32-bit "
     "(FP32) để giảm kích thước và tăng tốc, đánh đổi một phần độ chính xác.")

h2("2.3. Nhận diện đối tượng và đọc hiểu hình ảnh bằng AI")
para("(Lưu ý: mục này thay cho phần “nhận diện khuôn mặt” trong mẫu, để phù hợp với hai bài "
     "toán cốt lõi của đề tài.) Ứng dụng phải giải quyết hai bài toán nhận diện khác nhau "
     "nên dùng hai họ mô hình tương ứng.")
h3("2.3.1. Phát hiện đối tượng và họ mô hình YOLO (nhận diện tiền)")
para("Khác với phân loại ảnh (chỉ trả về một nhãn cho cả ảnh), phát hiện đối tượng (object "
     "detection) đồng thời xác định vị trí (khung bao) và nhãn của từng đối tượng. Điều này "
     "phù hợp để định vị và đếm nhiều tờ tiền trong khung hình. YOLO (You Only Look Once) là "
     "họ mô hình phát hiện một giai đoạn, dự đoán khung bao và xác suất lớp trong một lần lan "
     "truyền nên rất nhanh. YOLOv10 cải tiến bằng cách bỏ bước hậu xử lý NMS nhờ gán nhãn "
     "kép, giảm độ trễ; biến thể YOLOv10n (nano) là phiên bản nhẹ nhất, hợp với điện thoại. "
     "Đề tài chọn YOLOv10n cho mô-đun nhận diện tiền vì xử lý tốt tờ tiền bị gấp, vò, che — "
     "tình huống phổ biến khi người khiếm thị cầm tiền.")
para("Về cấu trúc, một mô hình YOLO điển hình gồm ba phần: phần trích đặc trưng (backbone) "
     "rút các đặc trưng từ ảnh; phần tổng hợp đặc trưng đa tỉ lệ (neck) giúp phát hiện đối "
     "tượng ở nhiều kích thước; và phần đầu ra (head) dự đoán khung bao, lớp và độ tin cậy. "
     "YOLOv10 sử dụng thiết kế không cần neo (anchor-free) và chiến lược gán nhãn kép giúp "
     "loại bỏ bước NMS hậu xử lý, nhờ đó suy luận nhanh và ổn định hơn — rất phù hợp với yêu "
     "cầu thời gian thực trên điện thoại của đề tài.")
table(
    ["Tiêu chí", "Phân loại ảnh (MobileNetV3)", "Phát hiện đối tượng (YOLOv10n)"],
    [
        ["Đầu ra", "Một nhãn / ảnh", "Nhiều khung bao + nhãn + độ tin cậy"],
        ["Đếm nhiều tờ", "Không hỗ trợ trực tiếp", "Hỗ trợ tự nhiên"],
        ["Tiền bị gấp/che", "Dễ nhầm khi nền phức tạp", "Định vị tốt hơn nhờ học vùng"],
        ["Vai trò trong đề tài", "Giải pháp ban đầu (v0.3)", "Giải pháp hoàn thiện (v0.4+)"],
    ],
    widths=[3.8, 5.8, 5.9],
    caption_text="Bảng 2.3. So sánh hai cách tiếp cận cho mô-đun nhận diện tiền",
)
h3("2.3.2. Mô hình ngôn ngữ – thị giác VLM (đọc thực đơn)")
para("Mô hình ngôn ngữ lớn (Large Language Model – LLM) dựa trên kiến trúc Transformer, huấn "
     "luyện trên khối lượng văn bản khổng lồ, có khả năng hiểu và sinh ngôn ngữ tự nhiên. Mô "
     "hình ngôn ngữ – thị giác (Vision-Language Model – VLM) bổ sung bộ mã hóa hình ảnh, cho "
     "phép mô hình “nhìn” ảnh và sinh văn bản theo nội dung ảnh. Đề tài dùng VLM Llama 4 "
     "Scout (Meta) qua dịch vụ Groq để đọc menu: thay vì OCR rồi phân tích cú pháp (vốn làm "
     "“phẳng” bố cục nhiều cột, khó ghép đúng món với giá), VLM nhìn trực tiếp ảnh, hiểu quan "
     "hệ không gian giữa tên món và giá, trả về kết quả JSON có cấu trúc.")
h3("2.3.3. OCR và chuyển văn bản thành giọng nói (TTS)")
para("OCR (Optical Character Recognition) chuyển hình ảnh chứa chữ thành văn bản; thư viện "
     "Google ML Kit Text Recognition v2 hỗ trợ tốt tiếng Việt có dấu và từng được dùng ở giai "
     "đoạn đầu. TTS (Text-to-Speech) là kênh đầu ra chính cho người khiếm thị; đề tài dùng "
     "FPT.AI TTS (đám mây, giọng tự nhiên) làm ưu tiên và Android TextToSpeech làm dự phòng "
     "ngoại tuyến. Chiều ngược lại, nhận dạng giọng nói (SpeechRecognizer của Android) hiện "
     "thực điều khiển bằng lệnh tiếng Việt.")

h2("2.4. Công nghệ và thư viện sử dụng")
table(
    ["Thành phần", "Công nghệ sử dụng"],
    [
        ["Ngôn ngữ & nền tảng", "Kotlin, Android (minSdk 26, target/compileSdk 36)"],
        ["Giao diện", "Jetpack Compose, Material 3, Navigation Compose"],
        ["Camera", "CameraX (Preview + ImageAnalysis)"],
        ["AI on-device", "TensorFlow Lite + TFLite Support"],
        ["OCR", "Google ML Kit Text Recognition v2"],
        ["AI đám mây", "Groq Cloud – Llama 4 Scout (Vision)"],
        ["TTS", "FPT.AI TTS + Android TextToSpeech"],
        ["Giọng nói", "Android SpeechRecognizer"],
        ["Tiêm phụ thuộc (DI)", "Hilt (Dagger)"],
        ["Bất đồng bộ", "Kotlin Coroutines + Flow"],
        ["Lưu trữ", "Room (lịch sử), DataStore (cài đặt)"],
        ["Huấn luyện mô hình", "Google Colab, PyTorch/Ultralytics → ONNX → TFLite"],
    ],
    widths=[5.0, 9.5],
    caption_text="Bảng 2.1. Tổng hợp công nghệ và thư viện sử dụng",
)

h2("2.5. Quy trình huấn luyện và đánh giá mô hình học sâu")
h3("2.5.1. Hàm mất mát và tối ưu hóa")
para("Huấn luyện một mạng nơ-ron là quá trình điều chỉnh các trọng số sao cho dự đoán của "
     "mô hình gần với nhãn thật nhất. Mức “sai” được đo bằng hàm mất mát (loss function): "
     "với bài toán phân loại thường dùng entropy chéo (cross-entropy), với bài toán phát "
     "hiện đối tượng dùng tổ hợp mất mát phân loại, mất mát hồi quy khung bao và mất mát độ "
     "tin cậy. Thuật toán lan truyền ngược (backpropagation) tính đạo hàm của hàm mất mát "
     "theo từng trọng số; thuật toán tối ưu (ví dụ SGD, Adam, AdamW) dùng đạo hàm đó để cập "
     "nhật trọng số theo hướng giảm mất mát. Mỗi lần mô hình duyệt qua toàn bộ tập huấn "
     "luyện gọi là một epoch; tốc độ học (learning rate) quyết định bước cập nhật.")
h3("2.5.2. Quá khớp và các kỹ thuật chống quá khớp")
para("Quá khớp (overfitting) xảy ra khi mô hình “học thuộc” dữ liệu huấn luyện nhưng kém "
     "tổng quát trên dữ liệu mới. Các kỹ thuật phổ biến để hạn chế gồm: tăng cường dữ liệu "
     "(data augmentation – xoay, lật, đổi sáng, cắt ngẫu nhiên), dừng sớm (early stopping), "
     "điều chuẩn (regularization), dropout và học chuyển giao. Trong đề tài, tăng cường dữ "
     "liệu đặc biệt quan trọng do số ảnh tiền thật còn hạn chế.")
h3("2.5.3. Các độ đo đánh giá")
para("Để đánh giá khách quan, đề tài sử dụng các độ đo chuẩn của bài toán phân loại và phát "
     "hiện đối tượng:")
table(
    ["Độ đo", "Ý nghĩa"],
    [
        ["Accuracy (độ chính xác)", "Tỷ lệ mẫu được dự đoán đúng trên tổng số mẫu"],
        ["Precision (độ chuẩn xác)", "Trong các dự đoán dương, bao nhiêu phần là đúng"],
        ["Recall (độ bao phủ)", "Trong các mẫu dương thật, mô hình tìm ra được bao nhiêu"],
        ["F1-score", "Trung bình điều hòa của Precision và Recall"],
        ["IoU", "Tỷ lệ diện tích giao trên hợp giữa khung bao dự đoán và khung bao thật"],
        ["mAP@0.5", "Độ chính xác trung bình khi coi dự đoán đúng nếu IoU ≥ 0,5"],
        ["Confusion matrix", "Ma trận nhầm lẫn giữa các lớp, giúp tìm lớp hay bị nhầm"],
    ],
    widths=[5.0, 9.5],
    caption_text="Bảng 2.2. Các độ đo đánh giá mô hình sử dụng trong đề tài",
)

h2("2.6. Triển khai mô hình học sâu trên thiết bị di động")
para("Mô hình sau khi huấn luyện (thường ở định dạng PyTorch) cần được chuyển đổi để chạy "
     "hiệu quả trên điện thoại. Quy trình điển hình là PyTorch → ONNX → TensorFlow Lite. "
     "ONNX là định dạng trung gian mở giúp trao đổi mô hình giữa các khung. TensorFlow Lite "
     "(TFLite) là phiên bản gọn nhẹ của TensorFlow dành cho thiết bị biên (edge), hỗ trợ:")
bullet("Lượng tử hóa (quantization): biểu diễn trọng số bằng INT8 thay cho FP32, giảm kích "
       "thước tệp khoảng 4 lần và tăng tốc suy luận, đổi lại mất một phần độ chính xác.")
bullet("Bộ tăng tốc (delegate): tận dụng GPU, NPU hoặc NNAPI của thiết bị để tăng tốc.")
bullet("Ánh xạ bộ nhớ (memory-map): nạp mô hình trực tiếp từ tệp không nén, tăng tốc khởi tạo.")
para("Trong đề tài, mô hình YOLOv10n được xuất sang TFLite (FP32, tích hợp NMS) và nhúng "
     "vào thư mục assets của ứng dụng dưới dạng không nén để tăng tốc nạp.")

h2("2.7. Kiến trúc Transformer và cơ chế chú ý (Attention)")
para("Mô hình ngôn ngữ – thị giác dùng trong mô-đun đọc menu dựa trên kiến trúc Transformer. "
     "Khác với mạng hồi tiếp (RNN) xử lý tuần tự, Transformer dùng cơ chế tự chú ý "
     "(self-attention) cho phép mỗi phần tử trong chuỗi “nhìn” tới mọi phần tử khác và học "
     "trọng số quan hệ giữa chúng. Nhờ đó mô hình nắm bắt ngữ cảnh dài và song song hóa tốt. "
     "Trong VLM, ảnh được chia thành các mảnh nhỏ (patch), mã hóa thành chuỗi token thị giác "
     "rồi đưa vào Transformer cùng với token văn bản; cơ chế chú ý chéo (cross-attention) "
     "giúp mô hình liên kết vùng ảnh với từ ngữ, ví dụ liên kết một dòng tên món với con số "
     "giá nằm cùng hàng. Đây chính là lý do VLM ghép “món – giá” chính xác hơn cách OCR rồi "
     "phân tích cú pháp thuần túy.")

h2("2.8. Nguyên tắc thiết kế khả năng truy cập (Accessibility)")
para("Vì người dùng mục tiêu không nhìn thấy hoặc nhìn rất hạn chế, thiết kế khả năng truy "
     "cập là trọng tâm chứ không phải bổ sung. Đề tài tuân theo các nguyên tắc:")
bullet("Đa kênh phản hồi: mọi thao tác đều có phản hồi âm thanh (TTS) và/hoặc rung, không "
       "chỉ thị giác.")
bullet("Cử chỉ đơn giản, nhất quán: cùng một cử chỉ có ý nghĩa tương tự ở các màn hình để "
       "người dùng ghi nhớ một lần.")
bullet("Tương phản cao và vùng chạm lớn (≥ 64dp) cho người khiếm thị nhẹ còn nhìn được.")
bullet("Tương thích TalkBack (trình đọc màn hình của Android): cung cấp mô tả nội dung "
       "(content description), thứ tự focus hợp lý.")
bullet("Khoan dung lỗi: có cơ chế dừng khẩn (lắc máy / nút giảm âm), chống thao tác nhầm và "
       "chống đếm trùng.")
para("Các nguyên tắc này tham chiếu tinh thần của bộ hướng dẫn WCAG (Web Content "
     "Accessibility Guidelines) và tài liệu thiết kế khả năng truy cập của Android.")
caption("Hình 2.1. Minh họa kiến trúc CNN / YOLO / Transformer (chèn sơ đồ minh họa)")

h2("2.9. Các công nghệ nền tảng Android sử dụng")
para("Bên cạnh các mô hình AI, đề tài sử dụng một số thành phần nền tảng Android hiện đại:")
bullet("bộ công cụ giao diện khai báo (declarative UI) mới của Android, "
       "giúp xây dựng giao diện ngắn gọn và phản ứng theo trạng thái.",
       bold_prefix="Jetpack Compose: ")
bullet("thư viện camera trừu tượng hóa khác biệt phần cứng, cung cấp luồng xem trước "
       "(Preview) và phân tích khung hình (ImageAnalysis) phục vụ nhận diện thời gian thực.",
       bold_prefix="CameraX: ")
bullet("khung tiêm phụ thuộc giúp quản lý vòng đời đối tượng, tách rời các mô-đun và thuận "
       "tiện thay thế (ví dụ thay mô hình nhận diện tiền).", bold_prefix="Hilt: ")
bullet("thư viện cơ sở dữ liệu trừu tượng trên SQLite, lưu lịch sử quét an toàn theo kiểu.",
       bold_prefix="Room: ")
bullet("giải pháp lưu cấu hình bất đồng bộ, thay cho SharedPreferences cũ.",
       bold_prefix="DataStore: ")
bullet("mô hình lập trình bất đồng bộ gọn nhẹ, kết hợp luồng dữ liệu phản ứng "
       "(Flow) để cập nhật giao diện theo trạng thái.", bold_prefix="Coroutines & Flow: ")
para("Việc lựa chọn bộ công nghệ Jetpack hiện đại giúp mã nguồn ngắn gọn, dễ kiểm thử và "
     "bảo trì, đồng thời tận dụng tốt các tính năng khả năng truy cập sẵn có của nền tảng.")

# ============================================================
#  III. PHÂN TÍCH HỆ THỐNG
# ============================================================
h1("III. Phân tích hệ thống")

h2("3.1. Kiến trúc tổng thể")
para("Ứng dụng tổ chức theo mẫu kiến trúc MVVM (Model – View – ViewModel) kết hợp tiêm phụ "
     "thuộc bằng Hilt. Tầng View (Jetpack Compose) chỉ hiển thị trạng thái và phát sự kiện; "
     "tầng ViewModel giữ logic nghiệp vụ và trạng thái dạng StateFlow; tầng dữ liệu gồm các "
     "mô-đun AI, TTS và lưu trữ. Mô hình nhận diện tiền chạy on-device, còn mô hình đọc menu "
     "gọi qua đám mây.")
code(
"            ┌───────────────────────────────────────────┐\n"
"            │           Tầng Giao diện (Compose)         │\n"
"            │   MoneyScreen · MenuScreen · Settings...   │\n"
"            └───────────────┬───────────────────────────┘\n"
"                            │ sự kiện / StateFlow\n"
"            ┌───────────────▼───────────────────────────┐\n"
"            │            Tầng ViewModel (MVVM)           │\n"
"            │   MoneyViewModel · MenuViewModel · ...     │\n"
"            └───┬───────────┬───────────────┬───────────┘\n"
"                │           │               │\n"
"      ┌─────────▼──┐  ┌─────▼──────┐  ┌─────▼──────────┐\n"
"      │ Mô-đun ML  │  │ TtsManager │  │  Lưu trữ        │\n"
"      │ YOLOv10n   │  │ FPT+Android│  │  Room/DataStore │\n"
"      │ (on-device)│  └────────────┘  └────────────────┘\n"
"      │ Groq VLM   │\n"
"      │ (cloud)    │\n"
"      └────────────┘\n"
)
caption("Hình 3.1. Sơ đồ kiến trúc tổng thể của ứng dụng Mắt AI")

h2("3.2. Chức năng hệ thống")
table(
    ["Mã", "Chức năng", "Mô tả"],
    [
        ["CN-01", "Nhận diện tiền", "Nhận diện mệnh giá theo thời gian thực, đọc to mệnh "
         "giá, đếm cộng dồn nhiều tờ và đọc tổng."],
        ["CN-02", "Đọc thực đơn", "Chụp ảnh menu, trích xuất danh sách món + giá, đọc lần "
         "lượt, chọn món và tính tổng tiền."],
        ["CN-03", "Phản hồi giọng nói", "Mọi kết quả và hướng dẫn đều đọc bằng TTS tiếng Việt."],
        ["CN-04", "Điều khiển bằng cử chỉ", "Chạm đôi, vuốt 4 hướng, giữ lâu, lắc máy."],
        ["CN-05", "Lệnh giọng nói", "Mở chức năng, đọc lại, dừng, đổi tốc độ bằng lời nói."],
        ["CN-06", "Lịch sử", "Lưu và phát lại 20 lần quét gần nhất."],
        ["CN-07", "Cài đặt", "Tùy chỉnh tốc độ đọc, giọng, tương phản, rung, tự động chụp."],
    ],
    widths=[1.6, 3.4, 9.5],
    caption_text="Bảng 3.1. Danh sách chức năng hệ thống",
)
para("Tác nhân duy nhất là Người dùng khiếm thị. Các ca sử dụng chính: đếm tiền, đọc menu, "
     "ra lệnh giọng nói, xem lịch sử, cấu hình.")
caption("Hình 3.2. Sơ đồ ca sử dụng tổng quát (chèn sơ đồ Use Case tại đây)")

h2("3.3. Luồng xử lý")
h3("3.3.1. Luồng nhận diện tiền")
code(
"Khung hình CameraX (ImageAnalysis, liên tục)\n"
"   → FrameQualityAnalyzer: kiểm tra độ sáng + độ nét\n"
"   → Yolov10MoneyDetector: resize 640×640, chuẩn hóa /255\n"
"   → lọc độ tin cậy ≥ 0,70 → ánh xạ lớp → mệnh giá VND\n"
"   → MoneyViewModel: ổn định khung → đọc TTS mệnh giá\n"
"   → người dùng vuốt xuống: cộng vào tổng → đọc tổng\n"
)
caption("Hình 3.3. Luồng xử lý nhận diện và đếm tiền")
h3("3.3.2. Luồng đọc thực đơn")
code(
"Người dùng hướng camera vào menu → tự chụp (đếm lùi 3-2-1)\n"
"   → thu nhỏ ảnh ≤ 1024px, nén JPEG, mã hóa base64\n"
"   → gửi Groq VLM (Llama 4 Scout) kèm prompt\n"
"   → nhận JSON {danh sách món, giá} → chuẩn hóa giá\n"
"   → TTS đọc lần lượt; vuốt phải/trái duyệt món; vuốt xuống chọn\n"
)
caption("Hình 3.4. Luồng xử lý đọc thực đơn")

h2("3.4. Yêu cầu phi chức năng")
table(
    ["Nhóm", "Yêu cầu"],
    [
        ["Khả năng truy cập", "Dùng được hoàn toàn khi không nhìn màn hình; tương thích "
         "TalkBack; tương phản cao; vùng chạm lớn."],
        ["Hiệu năng", "Nhận diện tiền theo thời gian thực (độ trễ mục tiêu < 200ms/khung); "
         "thời gian tới kết quả < 10 giây."],
        ["Độ tin cậy", "Có cơ chế dự phòng khi mô hình lỗi hoặc mất mạng; không treo ứng dụng."],
        ["Tính khả chuyển", "Chạy trên Android 8.0 trở lên; không yêu cầu phần cứng đặc biệt."],
        ["Bảo mật", "Khóa API lưu ở local.properties, không đưa vào mã nguồn công khai."],
        ["Bảo trì", "Kiến trúc tách lớp (MVVM), cô lập mô-đun AI để dễ thay thế mô hình."],
    ],
    widths=[3.5, 11.0],
    caption_text="Bảng 3.2. Yêu cầu phi chức năng",
)

h2("3.5. Đặc tả chi tiết ca sử dụng")
para("Hệ thống có một tác nhân chính là Người dùng khiếm thị. Dưới đây đặc tả hai ca sử "
     "dụng quan trọng nhất.")
h3("3.5.1. Ca sử dụng “Đếm tiền”")
table(
    ["Thành phần", "Nội dung"],
    [
        ["Mã / Tên", "UC-01 / Đếm tiền"],
        ["Tác nhân", "Người dùng khiếm thị"],
        ["Tiền điều kiện", "Đã cấp quyền camera; mô hình nhận diện sẵn sàng"],
        ["Luồng chính",
         "1) Người dùng mở chế độ đếm tiền; 2) Hệ thống bật camera và đọc hướng dẫn; "
         "3) Người dùng đưa tờ tiền vào khung; 4) Hệ thống nhận diện và đọc mệnh giá; "
         "5) Người dùng vuốt xuống để cộng tờ vào tổng; 6) Hệ thống đọc tổng; "
         "7) Lặp lại cho các tờ tiếp theo."],
        ["Luồng phụ", "Nếu không nhận diện được: hệ thống nhắc đưa tờ tiền vào giữa; nếu "
         "vuốt xuống khi chưa thấy tờ nào: thông báo phù hợp."],
        ["Hậu điều kiện", "Tổng số tiền được cập nhật; có thể lưu vào lịch sử khi chuyển chế độ."],
    ],
    widths=[3.2, 11.3],
    caption_text="Bảng 3.3. Đặc tả ca sử dụng Đếm tiền",
)
h3("3.5.2. Ca sử dụng “Đọc thực đơn”")
table(
    ["Thành phần", "Nội dung"],
    [
        ["Mã / Tên", "UC-02 / Đọc thực đơn"],
        ["Tác nhân", "Người dùng khiếm thị"],
        ["Tiền điều kiện", "Đã cấp quyền camera; có kết nối mạng; đã cấu hình khóa dịch vụ AI"],
        ["Luồng chính",
         "1) Người dùng mở chế độ đọc menu; 2) Hệ thống bật camera và đếm lùi bằng âm thanh; "
         "3) Hệ thống tự chụp ảnh menu; 4) Gửi ảnh tới mô hình ngôn ngữ – thị giác; "
         "5) Nhận danh sách món + giá; 6) Đọc lần lượt; 7) Người dùng vuốt phải/trái duyệt, "
         "vuốt xuống chọn món; 8) Hệ thống đọc tổng tiền các món đã chọn."],
        ["Luồng phụ", "Mất mạng / lỗi dịch vụ: thông báo và cho thử lại; menu không có món: "
         "thông báo và cho chụp lại."],
        ["Hậu điều kiện", "Danh sách món được đọc; lần quét được lưu vào lịch sử."],
    ],
    widths=[3.2, 11.3],
    caption_text="Bảng 3.4. Đặc tả ca sử dụng Đọc thực đơn",
)

h2("3.6. Biểu đồ tuần tự")
para("Biểu đồ tuần tự (sequence diagram) mô tả thứ tự tương tác giữa các thành phần khi "
     "thực hiện một chức năng. Với chức năng đếm tiền, chuỗi tương tác chính như sau:")
code(
"Người dùng → MoneyScreen : đưa tờ tiền / vuốt xuống\n"
"MoneyScreen → CameraX : lấy khung hình liên tục\n"
"CameraX → Yolov10MoneyDetector : khung hình\n"
"Yolov10MoneyDetector → MoneyViewModel : MoneyResult (mệnh giá, độ tin cậy)\n"
"MoneyViewModel → TtsManager : đọc mệnh giá / tổng\n"
"TtsManager → Người dùng : phát giọng nói\n"
)
caption("Hình 3.5. Biểu đồ tuần tự chức năng đếm tiền (chèn sơ đồ UML nếu cần)")
caption("Hình 3.6. Biểu đồ tuần tự chức năng đọc thực đơn (chèn sơ đồ UML nếu cần)")

h2("3.7. Thiết kế cơ sở dữ liệu cục bộ")
para("Lịch sử quét được lưu bằng Room (SQLite) qua thực thể ScanHistoryEntity; cài đặt người "
     "dùng lưu bằng Jetpack DataStore (Preferences).")
table(
    ["Trường", "Kiểu", "Ý nghĩa"],
    [
        ["id", "Long (PK, tự tăng)", "Khóa chính"],
        ["type", "ScanType (MONEY/MENU)", "Loại lần quét"],
        ["content", "String", "Nội dung đã đọc bằng TTS"],
        ["timestamp", "Long", "Thời điểm quét (epoch millis)"],
    ],
    widths=[3.0, 4.8, 6.7],
    caption_text="Bảng 3.5. Cấu trúc bảng lịch sử quét (ScanHistoryEntity)",
)
table(
    ["Cài đặt", "Kiểu", "Mặc định"],
    [
        ["Tốc độ đọc", "Float (0.5–2.0)", "1.0"],
        ["Giọng đọc", "Enum (Tự động / FPT / Android)", "Tự động"],
        ["Độ tương phản", "Enum (Hệ thống / Đen-vàng / Trắng-đen)", "Hệ thống"],
        ["Rung phản hồi", "Boolean", "Bật"],
        ["Tự động chụp tiền", "Boolean", "Tắt"],
        ["Điều khiển giọng nói", "Boolean", "Bật"],
    ],
    widths=[4.5, 6.0, 4.0],
    caption_text="Bảng 3.6. Các tùy chọn cài đặt (DataStore)",
)

h2("3.8. Thiết kế các lớp chính")
table(
    ["Lớp / Thành phần", "Trách nhiệm"],
    [
        ["MainActivity", "Điểm vào ứng dụng, host Compose, xử lý quyền và phím cứng"],
        ["AppNavHost", "Điều hướng giữa các màn hình bằng Navigation Compose"],
        ["MoneyViewModel", "Logic đếm tiền, ổn định khung, chống đếm trùng, đọc tổng"],
        ["MenuViewModel", "Logic đọc menu, đếm lùi, duyệt/chọn món, tính tổng"],
        ["MoneyClassifier (interface)", "Hợp đồng nhận diện tiền (cho phép thay mô hình)"],
        ["Yolov10MoneyDetector", "Nhận diện tiền bằng YOLOv10n trên TFLite"],
        ["GroqMenuAnalyzer", "Gọi mô hình ngôn ngữ – thị giác, phân tích JSON món + giá"],
        ["TtsManager / RoutedTtsEngine", "Định tuyến TTS giữa FPT.AI và Android"],
        ["VoiceCommandService", "Nhận và ánh xạ lệnh giọng nói tiếng Việt"],
        ["FrameQualityAnalyzer", "Đánh giá độ sáng/độ nét để tự động chụp"],
        ["HistoryRepository", "Truy xuất và ghi lịch sử quét (Room)"],
        ["PreferencesRepository", "Đọc/ghi cài đặt người dùng (DataStore)"],
    ],
    widths=[5.2, 9.3],
    caption_text="Bảng 3.7. Các lớp/thành phần chính và trách nhiệm",
)
caption("Hình 3.7. Sơ đồ lớp tổng quát (chèn sơ đồ class UML nếu cần)")

# ============================================================
#  IV. TRIỂN KHAI VÀ THỰC NGHIỆM
# ============================================================
h1("IV. Triển khai và thực nghiệm")

h2("4.1. Chuẩn bị")
para("Môi trường phát triển: Android Studio, JDK 17, Gradle 8.9, thiết bị Android thật (API "
     "26 trở lên) có camera sau. Mô hình AI được huấn luyện trên Google Colab (GPU T4). Khóa "
     "API (Groq, FPT.AI) lưu trong tệp local.properties (không đưa vào mã nguồn) và được nạp "
     "vào BuildConfig khi biên dịch.")
para("Một đặc điểm đáng chú ý của đề tài là kiến trúc được điều chỉnh qua nhiều phiên bản dựa "
     "trên kết quả thực nghiệm:")
table(
    ["Phiên bản", "Mô-đun tiền", "Mô-đun menu", "Lý do thay đổi"],
    [
        ["v0.3", "MobileNetV3 (phân loại)", "ML Kit OCR + regex", "Giải pháp khởi đầu"],
        ["v0.4", "YOLOv10n (phát hiện)", "Gemini 1.5 Flash (VLM)",
         "OCR mất bố cục cột; phân loại khó với tiền gấp/che"],
        ["v0.5–0.6", "YOLOv10n", "Groq Llama 4 Scout (Vision)",
         "VLM nhìn ảnh trực tiếp ghép “món–giá” chính xác hơn"],
        ["v0.7–0.9", "YOLOv10n + đếm cộng dồn", "Điều hướng + chọn món",
         "Hoàn thiện trải nghiệm cho người khiếm thị"],
    ],
    widths=[2.4, 4.0, 4.1, 4.0],
    caption_text="Bảng 4.1. Các mốc tiến hóa kiến trúc của Mắt AI",
)

h2("4.2. Mô-đun nhận diện tiền")
para("(Mục này thay cho phần “đăng ký khuôn mặt” trong mẫu.) Mô-đun nhận diện tiền là một "
     "trong hai chức năng lõi của ứng dụng.")
h3("4.2.1. Dữ liệu và huấn luyện")
para("Bộ dữ liệu chính lấy từ Roboflow Universe, gồm ảnh tiền VND ở nhiều trạng thái thực "
     "tế (phẳng, gấp, vò nhàu, bị che một phần), gán nhãn theo định dạng phát hiện đối tượng "
     "cho 9 mệnh giá. Quy trình huấn luyện đóng gói trong notebook Colab.")
table(
    ["Tham số", "Giá trị"],
    [
        ["Kiến trúc", "YOLOv10n (nano)"],
        ["Trọng số khởi tạo", "yolov10n.pt (tiền huấn luyện COCO)"],
        ["Số epoch", "100"],
        ["Kích thước ảnh đầu vào", "640 × 640"],
        ["Batch size", "Tự động (auto-batch)"],
        ["Số lớp", "9 mệnh giá VND"],
        ["Định dạng xuất", "TensorFlow Lite, FP32, tích hợp NMS"],
    ],
    widths=[6.0, 8.5],
    caption_text="Bảng 4.2. Cấu hình huấn luyện mô hình YOLOv10n",
)
para("Dữ liệu được chia thành ba tập theo tỷ lệ phổ biến: huấn luyện (train), kiểm định "
     "(validation) và kiểm tra (test). Để tăng tính tổng quát, quá trình huấn luyện áp dụng "
     "tăng cường dữ liệu (data augmentation) như:")
bullet("Biến đổi hình học: xoay, lật, dịch, co giãn, cắt ngẫu nhiên — mô phỏng các góc cầm "
       "tiền khác nhau.")
bullet("Biến đổi màu/sáng: đổi độ sáng, tương phản, sắc độ — mô phỏng điều kiện ánh sáng "
       "phòng, đèn vàng, thiếu sáng.")
bullet("Làm mờ và thêm nhiễu — mô phỏng ảnh rung tay hoặc lấy nét chưa chuẩn của người "
       "khiếm thị.")
para("Đề tài loại bỏ ba mệnh giá 100đ, 200đ, 500đ vì gần như không còn lưu hành, đồng thời "
     "chú trọng cân bằng số lượng mẫu giữa 9 mệnh giá còn lại để mô hình không thiên lệch. "
     "Việc bổ sung ảnh “không phải tiền” giúp giảm báo nhầm khi camera chĩa vào vật phẳng "
     "khác như danh thiếp, giấy.")
h3("4.2.2. Suy luận trên thiết bị và cơ chế dự phòng")
para("Lớp Yolov10MoneyDetector nạp mô hình bằng ánh xạ bộ nhớ (memory-map) để tăng tốc, "
     "chuẩn hóa khung hình về 640×640 và giá trị điểm ảnh về [0,1], lấy phát hiện có độ tin "
     "cậy cao nhất, lọc theo ngưỡng 0,70 rồi ánh xạ sang mệnh giá. Mô-đun thiết kế dự phòng "
     "nhiều tầng để luôn chạy được: YOLOv10n (chính) → MobileNetV3 (dự phòng) → bộ giả lập "
     "(kiểm thử giao diện). Mọi bộ nhận diện đều cài giao diện MoneyClassifier nên có thể "
     "thay mô hình mà không sửa tầng trên.")
h3("4.2.3. Đếm cộng dồn và đọc số thành chữ")
para("MoneyViewModel hiện thực logic “đếm cộng dồn”: khi một tờ tiền được nhận diện ổn định, "
     "ứng dụng đọc mệnh giá nhưng chưa cộng; người dùng vuốt xuống để xác nhận cộng vào tổng; "
     "sau đó hệ thống yêu cầu nhấc tờ tiền ra trước khi đếm tờ kế tiếp để tránh đếm trùng. "
     "Hàm NumberToVietnamese tự xây xử lý đầy đủ các đặc thù tiếng Việt: “mốt” (21), “lăm” "
     "(25), “linh” (101), nhóm số 0 ở giữa (1.000.001 → một triệu không trăm linh một).")

h2("4.3. Mô-đun đọc thực đơn")
para("(Mục này thay cho phần “nhận diện và điểm danh” trong mẫu.) MenuViewModel điều phối "
     "luồng đọc menu. Sau khi chụp, ứng dụng gửi thẳng ảnh tới Groq Vision (Llama 4 Scout) "
     "— bỏ qua OCR vì OCR làm phẳng bố cục cột khiến ghép món–giá kém chính xác. Ảnh được thu "
     "nhỏ về tối đa 1024px, nén JPEG, mã hóa base64 và gửi kèm prompt yêu cầu trả về JSON "
     "danh sách {tên món, giá} đã chuẩn hóa.")
code(
'{"items":[{"name":"Phở Bò","price":"50000"},\n'
'          {"name":"Bún Chả","price":"45000"}]}\n'
)
caption("Hình 4.1. Định dạng JSON do VLM trả về cho mô-đun đọc menu")
para("Lớp GroqMenuAnalyzer xử lý gọi API có cơ chế thử lại với độ trễ tăng dần cho lỗi tạm "
     "thời (HTTP 429/503), tách JSON khỏi văn bản thừa và chuẩn hóa giá (“50k”, “50.000”, "
     "“1.5tr” đều quy về số nguyên VND). Trải nghiệm đọc menu cho phép vuốt phải/trái duyệt "
     "món, vuốt xuống chọn món, đọc tổng tiền các món đã chọn — đối xứng với mô-đun đếm tiền "
     "để người dùng dễ ghi nhớ.")

h2("4.4. Kết quả thực nghiệm")
para("Ứng dụng đã chạy thông suốt hai luồng lõi trên thiết bị thật (phiên bản 0.9.0). Phần "
     "logic trọng yếu (đọc số tiền) được phủ kiểm thử đơn vị đầy đủ. Các bảng dưới là khuôn "
     "kết quả; các ô “[…]” cần điền số đo thực tế sau khi đánh giá cuối cùng.")
table(
    ["Chỉ số (nhận diện tiền)", "Mục tiêu", "Kết quả đo (điền sau)"],
    [
        ["mAP@0.5 (phát hiện)", "≥ 0,90", "[…]"],
        ["Độ chính xác mệnh giá (top-1)", "> 92%", "[…]"],
        ["Độ trễ suy luận / khung", "< 200 ms", "[…]"],
        ["Nhận diện tiền gấp/che", "Ổn định", "[…]"],
    ],
    widths=[6.0, 4.0, 4.5],
    caption_text="Bảng 4.3. Khuôn đánh giá mô-đun nhận diện tiền",
)
table(
    ["Chỉ số (đọc menu)", "Giá trị (điền sau)"],
    [
        ["Số ảnh menu đánh giá", "840"],
        ["Precision phát hiện món", "[…]"],
        ["Recall phát hiện món", "[…]"],
        ["F1 phát hiện món", "[…]"],
        ["Độ chính xác trích xuất giá", "[…]"],
    ],
    widths=[7.5, 7.0],
    caption_text="Bảng 4.4. Khuôn đánh giá mô-đun đọc thực đơn",
)
para("Phần kiểm thử đơn vị cho hàm đọc số thành chữ (NumberToVietnameseTest) đã chạy đạt toàn "
     "bộ, phủ các nhóm: 9 mệnh giá, hàng chục “mốt/lăm”, hàng trăm “linh”, nhóm 0 ở giữa, "
     "hàng triệu/tỷ, tổng đếm thực tế và các trường hợp biên.")

para("Về mặt định tính, qua quá trình thử nghiệm trên thiết bị thật, mô-đun nhận diện tiền "
     "hoạt động ổn định với tờ tiền phẳng và đủ sáng; với tờ gấp/vò/che, việc dùng mô hình "
     "phát hiện đối tượng cho kết quả tốt hơn rõ rệt so với phương án phân loại ảnh ban đầu. "
     "Cơ chế đếm cộng dồn kèm chống đếm trùng giúp người dùng đếm nhiều tờ liên tiếp một cách "
     "tự nhiên. Mô-đun đọc menu nhờ dùng mô hình ngôn ngữ – thị giác đã khắc phục được điểm "
     "yếu cố hữu của phương án OCR + luật là ghép sai món với giá ở menu nhiều cột.")
para("Các hạn chế quan sát được: độ chính xác giảm khi ánh sáng quá yếu hoặc tờ tiền bị che "
     "quá nửa; mô-đun đọc menu phụ thuộc chất lượng mạng và thời gian phản hồi của dịch vụ "
     "đám mây (thường 2–4 giây). Đây là cơ sở cho các hướng cải tiến ở Phần IX.")

h2("4.5. Một số đoạn mã nguồn minh họa")
h3("4.5.1. Mã nguồn suy luận nhận diện tiền (YOLOv10n)")
para("Đoạn mã rút gọn dưới đây minh họa cách lớp Yolov10MoneyDetector chuẩn hóa khung hình "
     "và lấy phát hiện có độ tin cậy cao nhất:")
code(
'override suspend fun classify(bitmap: Bitmap): MoneyResult =\n'
'  withContext(Dispatchers.Default) {\n'
'    val tensorImage = TensorImage(DataType.FLOAT32).apply { load(bitmap) }\n'
'    val input = processor.process(tensorImage)   // resize 640, /255\n'
'    interpreter.run(input.buffer, outputBuffer)\n'
'    val top = outputBuffer[0].firstOrNull { it[4] > 0f }\n'
'        ?: return@withContext MoneyResult.Unknown\n'
'    val conf = top[4]; val classId = top[5].toInt()\n'
'    if (conf < MIN_CONFIDENCE) return@withContext MoneyResult.Unknown\n'
'    val label = MONEY_LABELS[classId]\n'
'    MoneyResult.Recognized(label.denominationVnd, conf)\n'
'  }\n'
)
caption("Hình 4.2. Mã nguồn rút gọn suy luận nhận diện tiền (Kotlin)")

para("Hàm chuyển số tiền sang chữ tiếng Việt xử lý các trường hợp đặc thù:")
code(
'200_000.toVietnameseMoney()   // "Hai trăm nghìn đồng"\n'
'650_000.toVietnameseMoney()   // "Sáu trăm năm mươi nghìn đồng"\n'
'1_000_001.toVietnameseWords() // "một triệu không trăm linh một"\n'
'21.toVietnameseWords()        // "hai mươi mốt"  (mốt, không phải một)\n'
'25.toVietnameseWords()        // "hai mươi lăm"  (lăm, không phải năm)\n'
)
caption("Hình 4.3. Ví dụ kết quả hàm đọc số thành chữ tiếng Việt")

para("Đoạn mã cấu hình mô-đun ML theo cơ chế dự phòng nhiều tầng (Hilt):")
code(
'fun provideMoneyClassifier(ctx: Context): MoneyClassifier {\n'
'  runCatching { Yolov10MoneyDetector(ctx) }\n'
'      .onSuccess { return it }      // ưu tiên YOLOv10n\n'
'  runCatching { TfliteMoneyClassifier(ctx) }\n'
'      .onSuccess { return it }      // dự phòng MobileNetV3\n'
'  return FakeMoneyClassifier()      // giả lập khi chưa có mô hình\n'
'}\n'
)
caption("Hình 4.4. Cơ chế dự phòng nhiều tầng cho mô-đun nhận diện tiền")

h3("4.5.2. Mã nguồn gọi mô hình ngôn ngữ – thị giác")
para("Mô-đun đọc menu gửi ảnh kèm prompt yêu cầu trả về JSON; phần thân yêu cầu rút gọn:")
code(
'val body = JSONObject()\n'
'  .put("model", "meta-llama/llama-4-scout-17b-16e-instruct")\n'
'  .put("temperature", 0.1)\n'
'  .put("messages", /* text prompt + ảnh base64 */ )\n'
'// Prompt: "Trích xuất TẤT CẢ món ăn và giá, trả về JSON\n'
'//  {\\"items\\":[{\\"name\\":\\"...\\",\\"price\\":\\"50000\\"}]}"\n'
)
caption("Hình 4.5. Cấu trúc yêu cầu gửi tới mô hình ngôn ngữ – thị giác (rút gọn)")

h2("4.6. Mô-đun chuyển văn bản thành giọng nói")
para("RoutedTtsEngine đóng vai trò bộ định tuyến TTS thông minh: nếu người dùng chọn giọng "
     "FPT và có mạng + khóa API hợp lệ thì dùng FPT.AI cho giọng tự nhiên; ngược lại dùng "
     "Android TTS làm dự phòng ngoại tuyến. Cách này đảm bảo chất lượng giọng đọc mà vẫn "
     "hoạt động khi không có mạng. Hệ thống cũng kiểm tra trạng thái mạng thực tế (có kết "
     "nối và đã xác thực) trước khi quyết định gọi dịch vụ đám mây.")
table(
    ["Tình huống", "Engine được chọn"],
    [
        ["Chọn “Android mặc định”", "Android TTS (luôn dùng)"],
        ["Chọn FPT/Tự động + có mạng + có khóa", "FPT.AI TTS"],
        ["Chọn FPT/Tự động + mất mạng/không khóa", "Android TTS (dự phòng)"],
    ],
    widths=[7.5, 7.0],
    caption_text="Bảng 4.5. Logic định tuyến TTS",
)

h2("4.7. Tự động chụp dựa trên chất lượng khung hình")
para("FrameQualityAnalyzer phân tích nhanh từng khung hình CameraX để quyết định thời điểm "
     "tự chụp: ước lượng độ sáng (trung bình kênh Y) và độ nét (phương sai Laplacian trên "
     "mẫu con của khung hình). Khi cả hai điều kiện đạt ngưỡng trong nhiều khung liên tiếp, "
     "hệ thống coi khung hình đã ổn định và kích hoạt chụp — giúp người khiếm thị không phải "
     "canh nút bấm. Cơ chế lấy mẫu con (1/10 điểm ảnh) giúp tính toán đủ nhanh để chạy trên "
     "luồng phân tích của camera mà không gây giật.")

h2("4.8. Điều khiển bằng giọng nói và cử chỉ đặc biệt")
para("VoiceCommandService bao bọc SpeechRecognizer của Android, nhận lệnh tiếng Việt (vi-VN) "
     "và ánh xạ linh hoạt theo từ khóa (so khớp “chứa”) để bền với cách nói tự nhiên. Người "
     "dùng kích hoạt bằng cách giữ phím tăng âm lượng. Các lệnh hỗ trợ gồm: mở đọc tiền, mở "
     "đọc menu, đọc lại, dừng, nhanh hơn, chậm hơn, lịch sử, cài đặt, thoát.")
para("Ngoài ra, cảm biến gia tốc được dùng cho cử chỉ “lắc để dừng đọc” (panic stop): khi "
     "TTS đang đọc mà người dùng lắc mạnh điện thoại, hệ thống dừng ngay. Cơ chế có ngưỡng "
     "gia tốc và thời gian chống dội (debounce) để tránh kích hoạt nhầm khi đi lại.")
table(
    ["Lệnh giọng nói", "Hành động"],
    [
        ["“đọc tiền”, “tiền”", "Mở chế độ đếm tiền"],
        ["“đọc menu”, “thực đơn”", "Mở chế độ đọc thực đơn"],
        ["“đọc lại”, “lặp lại”", "Phát lại nội dung gần nhất"],
        ["“dừng”, “stop”", "Dừng đọc"],
        ["“nhanh hơn” / “chậm hơn”", "Tăng / giảm tốc độ đọc"],
        ["“lịch sử” / “cài đặt”", "Mở lịch sử / cài đặt"],
        ["“thoát”, “về”", "Quay lại màn hình chính"],
    ],
    widths=[5.5, 9.0],
    caption_text="Bảng 4.6. Danh sách lệnh giọng nói hỗ trợ",
)

h2("4.9. Kiểm thử phần mềm")
para("Đề tài áp dụng ba mức kiểm thử: (1) kiểm thử đơn vị (unit test) cho phần logic; (2) "
     "kiểm thử tích hợp/đầu cuối cho từng mô-đun trên thiết bị thật; (3) kiểm thử trải "
     "nghiệm bằng TalkBack và bịt mắt mô phỏng người khiếm thị.")
para("Phần logic chuyển số thành chữ tiếng Việt — vốn quan trọng vì sai một chữ sẽ khiến "
     "người dùng nghe nhầm số tiền — được phủ bằng bộ kiểm thử đơn vị gồm nhiều nhóm trường "
     "hợp, tất cả đều đạt:")
table(
    ["Nhóm test", "Ví dụ kiểm tra", "Kết quả"],
    [
        ["9 mệnh giá VND", "200.000 → “hai trăm nghìn”", "Đạt"],
        ["Hàng đơn vị 0–9", "9 → “chín”", "Đạt"],
        ["10–19 dùng “mười”", "15 → “mười lăm”", "Đạt"],
        ["20–99 “mốt/lăm”", "21 → “hai mươi mốt”; 25 → “hai mươi lăm”", "Đạt"],
        ["100–999 “linh”", "101 → “một trăm linh một”", "Đạt"],
        ["Nhóm 0 ở giữa", "1.000.001 → “một triệu không trăm linh một”", "Đạt"],
        ["Hàng triệu, tỷ", "1.500.000.000 → “một tỷ năm trăm triệu”", "Đạt"],
        ["Tổng đếm thực tế", "650.000 → “Sáu trăm năm mươi nghìn đồng”", "Đạt"],
        ["Biên / ngoại lệ", "Số âm và > 999 tỷ ném ngoại lệ", "Đạt"],
    ],
    widths=[3.6, 7.4, 3.5],
    caption_text="Bảng 4.7. Kết quả kiểm thử đơn vị mô-đun đọc số thành chữ",
)
para("Ngoài ra còn có bộ kiểm thử cho hàm phân tích/chuẩn hóa giá tiền trong menu. Phần kiểm "
     "thử trải nghiệm tập trung vào các tiêu chí: hoàn thành tác vụ mà không cần nhìn, thời "
     "gian tới kết quả, và độ rõ ràng của phản hồi âm thanh. Bảng dưới là khuôn đánh giá "
     "trải nghiệm; các ô “[…]” cần điền sau khi thử nghiệm với người dùng thật.")
table(
    ["Tiêu chí", "Mục tiêu", "Kết quả (điền sau)"],
    [
        ["Hoàn thành đếm tiền không cần nhìn", "100% người thử", "[…]"],
        ["Hoàn thành đọc menu không cần nhìn", "≥ 80% người thử", "[…]"],
        ["Thời gian tới kết quả", "< 10 giây", "[…]"],
        ["Điểm khả dụng SUS", "≥ 70", "[…]"],
    ],
    widths=[6.5, 4.0, 4.0],
    caption_text="Bảng 4.8. Khuôn đánh giá trải nghiệm người dùng",
)

# ============================================================
#  V. MÔ HÌNH TRIỂN KHAI
# ============================================================
h1("V. Mô hình triển khai (Deployment Model)")
para("Mắt AI là ứng dụng di động nên mô hình triển khai khác với hệ thống máy chủ truyền "
     "thống. Đề tài phân tích ba phương án triển khai theo mức độ phụ thuộc hạ tầng, từ chạy "
     "ngay trên máy đến mở rộng đám mây.")
caption("Hình 5.1. Mô hình triển khai thực tế: thiết bị (on-device) kết hợp dịch vụ đám mây")

h2("5.1. Triển khai cục bộ (Local / On-device)")
para("Đây là phương án triển khai chính của đề tài. Mô-đun nhận diện tiền (YOLOv10n, TFLite) "
     "chạy hoàn toàn trên điện thoại, không cần mạng. Toàn bộ giao diện, logic đếm tiền, TTS "
     "Android dự phòng, cơ sở dữ liệu lịch sử (Room) và cài đặt (DataStore) đều nằm cục bộ. "
     "Ưu điểm: hoạt động ngoại tuyến, độ trễ thấp, bảo mật dữ liệu người dùng (ảnh không rời "
     "khỏi máy đối với chức năng tiền). Việc cài đặt chỉ cần một tệp APK.")
table(
    ["Tiêu chí", "Triển khai cục bộ"],
    [
        ["Yêu cầu mạng", "Không (mô-đun tiền)"],
        ["Độ trễ", "Thấp (suy luận tại chỗ)"],
        ["Bảo mật dữ liệu", "Cao (xử lý ngay trên máy)"],
        ["Phù hợp với", "Người dùng cá nhân, dùng hằng ngày"],
    ],
    widths=[5.5, 9.0],
    caption_text="Bảng 5.1. Đặc điểm triển khai cục bộ",
)

h2("5.2. Triển khai mạng nội bộ (LAN)")
para("Phương án này hướng tới các tổ chức như Hội Người mù, trung tâm bảo trợ, nơi nhiều "
     "người dùng chung một hạ tầng và muốn dùng được cả khi không có Internet công cộng. Khi "
     "đó, mô-đun đọc menu (vốn cần mô hình ngôn ngữ – thị giác) có thể được phục vụ bởi một "
     "máy chủ suy luận đặt trong mạng nội bộ: triển khai một VLM mã nguồn mở (ví dụ Llama "
     "Vision/LLaVA) trên một máy trạm có GPU, các điện thoại trong cùng mạng LAN gọi tới máy "
     "chủ này thay vì gọi ra đám mây. Tương tự, có thể tự dựng máy chủ TTS tiếng Việt nội bộ.")
para("Ưu điểm: không phụ thuộc dịch vụ bên thứ ba, dữ liệu menu không ra ngoài tổ chức, chi "
     "phí vận hành ổn định. Nhược điểm: cần đầu tư phần cứng GPU và nhân lực quản trị; chỉ "
     "dùng được trong phạm vi phủ sóng của mạng nội bộ.")
caption("Hình 5.2. Mô hình triển khai LAN: điện thoại ↔ máy chủ suy luận nội bộ (chèn sơ đồ)")

h2("5.3. Triển khai đám mây (phát triển nâng cao)")
para("Ở quy mô lớn, mô-đun đọc menu và TTS có thể đặt sau một dịch vụ đám mây do nhóm phát "
     "triển vận hành, đứng giữa ứng dụng và các nhà cung cấp AI. Cách này cho phép: quản lý "
     "tập trung khóa API và hạn mức; cân bằng tải, tự co giãn theo số người dùng; ghi nhận "
     "thống kê ẩn danh để cải thiện mô hình; cập nhật mô hình mà không cần phát hành lại APK. "
     "Hiện tại đề tài đang dùng trực tiếp dịch vụ Groq cho mô-đun menu — đây chính là một "
     "hình thức triển khai đám mây ở mức cơ bản; bước nâng cao là bổ sung lớp dịch vụ trung "
     "gian (backend) của riêng ứng dụng.")
table(
    ["Phương án", "Mạng", "Chi phí hạ tầng", "Phù hợp"],
    [
        ["Cục bộ (on-device)", "Không bắt buộc", "Thấp", "Cá nhân, dùng thường ngày"],
        ["Mạng nội bộ (LAN)", "Nội bộ", "Trung bình (GPU)", "Tổ chức, trung tâm"],
        ["Đám mây", "Bắt buộc", "Theo lưu lượng", "Quy mô lớn, nhiều người dùng"],
    ],
    widths=[4.0, 3.0, 3.8, 3.7],
    caption_text="Bảng 5.2. So sánh ba mô hình triển khai",
)

# ============================================================
#  VI. GIAO DIỆN NGƯỜI DÙNG
# ============================================================
h1("VI. Giao diện người dùng (UI Mockup)")
para("Vì người dùng mục tiêu là người khiếm thị, giao diện được thiết kế tối giản, tương "
     "phản cao, nút lớn (≥ 96dp) và quan trọng nhất là vận hành chủ yếu qua âm thanh và cử "
     "chỉ. Giao diện đồ họa chủ yếu phục vụ người khiếm thị nhẹ còn nhìn được và người hỗ "
     "trợ. Ứng dụng có các màn hình chính sau:")
bullet("Màn hình chính: bốn nút lớn xếp dọc — Đọc tiền, Đọc menu, Lịch sử, Cài đặt; chạm "
       "một lần để chọn (đọc tên), chạm đôi để kích hoạt.", bold_prefix="Màn hình chính: ")
bullet("camera toàn màn hình, dải trạng thái ở dưới, đọc mệnh giá tờ đang thấy và tổng đã "
       "chọn.", bold_prefix="Màn hình đếm tiền: ")
bullet("camera với đếm lùi âm thanh, sau khi có kết quả hiển thị món hiện tại và danh sách "
       "đã chọn.", bold_prefix="Màn hình đọc menu: ")
bullet("danh sách 20 lần quét gần nhất, mỗi mục có nút nghe lại.",
       bold_prefix="Màn hình lịch sử: ")
bullet("tốc độ đọc, giọng đọc, độ tương phản, rung, tự động chụp, điều khiển bằng giọng nói.",
       bold_prefix="Màn hình cài đặt: ")
para("Bộ cử chỉ điều khiển được thống nhất giữa hai chế độ chính để người dùng chỉ cần ghi "
     "nhớ một lần:")
table(
    ["Cử chỉ", "Chế độ đếm tiền", "Chế độ đọc menu"],
    [
        ["Vuốt xuống", "Chọn (cộng) tờ tiền đang thấy", "Chọn món hiện tại"],
        ["Vuốt phải / trái", "—", "Món tiếp theo / món trước"],
        ["Vuốt lên", "Chuyển sang đọc menu", "Chuyển sang đếm tiền"],
        ["Chạm đôi", "Đọc tổng đã chọn", "Đọc lại danh sách đã chọn"],
        ["Giữ lâu", "Xóa hết, đếm lại", "Quét lại menu mới"],
        ["Lắc máy / Vol−", "Dừng đọc (panic stop)", "Dừng đọc (panic stop)"],
        ["Giữ Vol+", "Kích hoạt lệnh giọng nói", "Kích hoạt lệnh giọng nói"],
    ],
    widths=[3.2, 5.6, 5.7],
    caption_text="Bảng 6.1. Bộ cử chỉ điều khiển thống nhất",
)
caption("Hình 6.1. Mockup màn hình chính (chèn ảnh thiết kế/ảnh chụp màn hình)")
caption("Hình 6.2. Mockup màn hình đếm tiền và đọc menu (chèn ảnh)")
para("(Ghi chú: chèn ảnh mockup hoặc ảnh chụp màn hình thực tế của ứng dụng vào các vị trí "
     "trên trước khi nộp báo cáo.)", italic=True)

h2("6.1. Màn hình chính")
para("Màn hình chính gồm bốn nút lớn (≥ 96dp) xếp dọc: Đọc tiền, Đọc menu, Lịch sử, Cài "
     "đặt. Chạm một lần để chọn (TalkBack/ứng dụng đọc tên và mô tả nút), chạm đôi để kích "
     "hoạt. Khi mở ứng dụng, một lời chào bằng giọng nói hướng dẫn ngắn gọn cách dùng. Bố "
     "cục đơn giản, tương phản cao, tránh để người dùng bị lạc.")
h2("6.2. Màn hình đếm tiền")
para("Camera chiếm toàn màn hình, phía dưới có dải trạng thái nền tối hiển thị tờ tiền đang "
     "nhận diện và tổng đã chọn. Khi đưa tờ tiền vào, ứng dụng đọc mệnh giá; người dùng vuốt "
     "xuống để cộng vào tổng, chạm đôi để nghe tổng, giữ lâu để xóa và đếm lại, vuốt lên để "
     "chuyển sang đọc menu. Hệ thống có nhắc nhở bằng giọng nói khi để yên quá lâu.")
h2("6.3. Màn hình đọc thực đơn")
para("Sau khi mở, ứng dụng đếm lùi bằng âm thanh rồi tự chụp. Khi có kết quả, màn hình hiển "
     "thị món hiện tại và danh sách món đã chọn; người dùng vuốt phải/trái để duyệt món, vuốt "
     "xuống để chọn, chạm đôi để nghe lại danh sách đã chọn, giữ lâu để quét menu mới. Mỗi "
     "món được đọc kèm giá, ví dụ “Phở bò, giá năm mươi nghìn đồng”.")
h2("6.4. Màn hình lịch sử và cài đặt")
para("Màn hình lịch sử hiển thị 20 lần quét gần nhất (cả tiền và menu), mỗi mục có nút nghe "
     "lại và có thể xóa toàn bộ. Màn hình cài đặt cho phép điều chỉnh: tốc độ đọc (0,5×–2×), "
     "giọng đọc (Tự động/FPT/Android), độ tương phản (Hệ thống/Đen-vàng/Trắng-đen), bật/tắt "
     "rung, tự động chụp tiền và điều khiển bằng giọng nói. Mỗi thay đổi đều có xác nhận bằng "
     "giọng nói để người khiếm thị biết kết quả.")
caption("Hình 6.3. Ảnh chụp màn hình cài đặt và lịch sử (chèn ảnh)")
caption("Hình 6.4. Người dùng thao tác bằng cử chỉ khi không nhìn màn hình (chèn ảnh)")

# ============================================================
#  VII. ỨNG DỤNG THỰC TẾ & KHẢ NĂNG MỞ RỘNG
# ============================================================
h1("VII. Ứng dụng thực tế và khả năng mở rộng")

h2("7.1. Ứng dụng thực tế")
bullet("Giúp người khiếm thị tự kiểm tra mệnh giá tiền khi mua bán, nhận tiền thừa, hạn chế "
       "rủi ro bị nhầm hoặc lừa.")
bullet("Giúp người khiếm thị tự đọc thực đơn, chủ động gọi món tại quán ăn, nhà hàng.")
bullet("Là công cụ hỗ trợ học tập và sinh hoạt cho học sinh, người cao tuổi suy giảm thị lực.")
bullet("Có thể triển khai tại Hội Người mù, trung tâm bảo trợ xã hội như một công cụ trợ "
       "năng miễn phí.")

h2("7.2. Khả năng mở rộng")
numlist("Đưa mô-đun đọc menu chạy on-device bằng VLM gọn nhẹ để hoạt động ngoại tuyến.")
numlist("Bổ sung phát hiện tiền giả qua đặc điểm bảo an; mở rộng nhận diện tiền nước ngoài.")
numlist("Thêm mô tả khung cảnh xung quanh (đối tượng, chướng ngại) hỗ trợ di chuyển an toàn.")
numlist("Mở rộng đọc văn bản tổng quát (hóa đơn, biển hiệu, hướng dẫn thuốc).")
numlist("Hoàn thiện trợ lý hội thoại điều khiển rảnh tay hoàn toàn bằng giọng nói.")
numlist("Phát hành trên Google Play, hỗ trợ đa nền tảng (iOS) và đa ngôn ngữ.")
para("Về mặt xã hội, nếu được hoàn thiện và phổ biến, ứng dụng có thể trở thành một công cụ "
     "trợ năng thiết thực, góp phần thu hẹp khoảng cách số cho người khiếm thị và thúc đẩy "
     "tinh thần “công nghệ vì cộng đồng”. Mô hình mã nguồn và quy trình huấn luyện tái lập "
     "được cũng tạo điều kiện để các nhóm nghiên cứu khác kế thừa và mở rộng.")

# ============================================================
#  VIII. KHÓ KHĂN
# ============================================================
h1("VIII. Những khó khăn trong quá trình thực hiện")
bullet("Dữ liệu tiền VND đa dạng, thực tế (gấp, vò, che, thiếu sáng) khó thu thập và gán "
       "nhãn đủ phong phú; lớp “không phải tiền” khó cân bằng.", bold_prefix="Dữ liệu: ")
bullet("Bố cục menu rất đa dạng (nhiều cột, món xuống dòng, giá theo khoảng) khiến cách OCR "
       "+ phân tích cú pháp ban đầu thất bại, buộc phải chuyển sang mô hình ngôn ngữ – thị "
       "giác.", bold_prefix="Bài toán menu: ")
bullet("Đưa mô hình từ PyTorch sang TensorFlow Lite, đồng bộ thứ tự lớp giữa notebook và "
       "mã Kotlin, tối ưu tốc độ suy luận trên thiết bị thật.",
       bold_prefix="Chuyển đổi mô hình: ")
bullet("Thiết kế trải nghiệm “dùng được khi nhắm mắt” đòi hỏi tư duy khác biệt: phản hồi "
       "âm thanh tức thì, chống đếm trùng, tránh thao tác nhầm, kiểm thử bằng TalkBack.",
       bold_prefix="Trải nghiệm cho người khiếm thị: ")
bullet("Giọng đọc Android cho tiếng Việt thiếu tự nhiên; phải tích hợp FPT.AI và xây bộ "
       "định tuyến TTS có dự phòng.", bold_prefix="Chất lượng TTS: ")
bullet("Khó tiếp cận đủ số lượng người khiếm thị để kiểm thử trải nghiệm quy mô lớn trong "
       "thời gian làm đồ án.", bold_prefix="Kiểm thử người dùng: ")

para("Bảng dưới tổng hợp các rủi ro chính trong quá trình thực hiện và giải pháp tương ứng "
     "đã hoặc sẽ áp dụng:")
table(
    ["Rủi ro", "Giải pháp"],
    [
        ["Dữ liệu tiền dễ quá khớp", "Tăng cường dữ liệu mạnh, chụp đa dạng nhiều điều kiện"],
        ["OCR sai bố cục menu", "Chuyển sang mô hình ngôn ngữ – thị giác nhìn ảnh trực tiếp"],
        ["Dịch vụ đám mây hết hạn mức", "Cơ chế thử lại với độ trễ tăng dần; định hướng on-device"],
        ["Giọng TTS thiếu tự nhiên", "Tích hợp FPT.AI, định tuyến có dự phòng Android"],
        ["Suy luận chậm máy yếu", "Lượng tử hóa INT8, bộ tăng tốc GPU/NNAPI"],
        ["Khó tìm người khiếm thị thử", "Kiểm thử sơ bộ bằng TalkBack + bịt mắt mô phỏng"],
    ],
    widths=[5.5, 9.0],
    caption_text="Bảng 8.1. Rủi ro và giải pháp",
)

# ============================================================
#  IX. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN
# ============================================================
h1("IX. Kết luận và hướng phát triển")

h2("9.1. Kết luận")
para("Đề tài đã hoàn thành các mục tiêu đặt ra: xây dựng hoàn chỉnh ứng dụng Android “Mắt "
     "AI” (phiên bản 0.9.0) chạy được trên thiết bị thật; triển khai mô-đun nhận diện tiền "
     "VND theo thời gian thực bằng YOLOv10n on-device có đếm cộng dồn; triển khai mô-đun đọc "
     "thực đơn bằng mô hình ngôn ngữ – thị giác; thiết kế mô hình tương tác accessibility-"
     "first nhất quán; và xây dựng quy trình huấn luyện tái lập được kèm kiểm thử đơn vị cho "
     "logic trọng yếu. Kết quả cho thấy việc kết hợp AI on-device với AI đám mây, đặt trong "
     "thiết kế lấy người khiếm thị làm trung tâm, là hướng đi khả thi và hữu ích.")

h2("9.2. Hạn chế")
bullet("Mô-đun đọc menu phụ thuộc kết nối mạng và dịch vụ đám mây bên thứ ba.")
bullet("Số liệu đánh giá định lượng (độ chính xác, độ trễ, điểm khả dụng) cần đo và hoàn "
       "thiện với người khiếm thị thật.")
bullet("Bộ dữ liệu tiền còn hạn chế về đa dạng điều kiện thực tế.")
bullet("Chưa hỗ trợ phát hiện tiền giả và mô tả vật thể xung quanh.")

h2("9.3. Hướng phát triển")
numlist("Đưa mô-đun menu chạy on-device để hoạt động ngoại tuyến hoàn toàn.")
numlist("Mở rộng dữ liệu và bổ sung phát hiện tiền giả.")
numlist("Thêm mô tả khung cảnh xung quanh hỗ trợ di chuyển.")
numlist("Hoàn thiện điều khiển bằng giọng nói và trợ lý hội thoại.")
numlist("Tối ưu hiệu năng (INT8, GPU delegate) và phát hành lên Google Play.")
numlist("Tổ chức đánh giá người dùng quy mô lớn cùng Hội Người mù.")

# ============================================================
#  PHỤ LỤC
# ============================================================
h1("Phụ lục")

h2("Phụ lục A. Danh mục hình ảnh cần chèn và vị trí")
para("Bảng dưới liệt kê toàn bộ hình cần bổ sung trong báo cáo cùng gợi ý nội dung. Tại mỗi "
     "vị trí trong bài đã có sẵn dòng chú thích (caption) “Hình X.Y …”; chỉ cần nhấp ngay "
     "phía trên dòng chú thích đó và chèn ảnh (Insert → Picture).")
table(
    ["Hình", "Vị trí trong báo cáo", "Nội dung nên chèn"],
    [
        ["Hình 2.1", "Mục 2.8 (cuối phần II)", "Sơ đồ minh họa kiến trúc CNN / YOLO / "
         "Transformer (vẽ hoặc lấy từ tài liệu, ghi nguồn)"],
        ["Hình 3.1", "Mục 3.1", "Sơ đồ kiến trúc tổng thể (có thể thay sơ đồ ASCII bằng "
         "sơ đồ vẽ đẹp)"],
        ["Hình 3.2", "Mục 3.2", "Sơ đồ ca sử dụng (Use Case) tổng quát"],
        ["Hình 3.3 / 3.4", "Mục 3.3", "Sơ đồ luồng xử lý đếm tiền / đọc menu (flowchart)"],
        ["Hình 3.5 / 3.6", "Mục 3.6", "Biểu đồ tuần tự (sequence) hai chức năng chính"],
        ["Hình 3.7", "Mục 3.8", "Sơ đồ lớp (class diagram) tổng quát"],
        ["Hình 4.1", "Mục 4.3", "Ảnh chụp màn hình kết quả đọc menu (danh sách món + giá)"],
        ["Hình 4.2–4.5", "Mục 4.2–4.3", "Có thể giữ dạng mã, hoặc thay bằng ảnh chụp đoạn "
         "code trong Android Studio cho đẹp"],
        ["Hình 4.6 (thêm)", "Mục 4.2.1", "Ảnh mẫu dữ liệu tiền (tờ phẳng, gấp, vò, che)"],
        ["Hình 4.7 (thêm)", "Mục 4.4", "Biểu đồ loss/accuracy theo epoch + ma trận nhầm lẫn"],
        ["Hình 4.8 (thêm)", "Mục 4.4", "Ảnh kết quả phát hiện tiền (khung bao + nhãn)"],
        ["Hình 5.1", "Mục 5 (đầu phần)", "Đã chèn sẵn — mô hình triển khai on-device + cloud"],
        ["Hình 5.2", "Mục 5.2", "Sơ đồ triển khai LAN — tùy chọn (điện thoại ↔ máy chủ nội bộ)"],
        ["Hình 6.1", "Phần VI", "Ảnh chụp màn hình chính (4 nút lớn)"],
        ["Hình 6.2", "Phần VI", "Ảnh chụp màn hình đếm tiền và đọc menu"],
        ["Hình 6.3 (thêm)", "Phần VI", "Ảnh chụp màn hình cài đặt và lịch sử"],
        ["Hình 6.4 (thêm)", "Phần VI", "Ảnh người dùng thật/bịt mắt thao tác bằng cử chỉ"],
    ],
    widths=[2.6, 4.4, 7.5],
    caption_text="Bảng A.1. Danh mục hình ảnh cần chèn",
)
para("Mẹo chèn ảnh chuyên nghiệp trong Word:", bold=True, space_before=8)
bullet("Chèn ảnh xong, bôi đen ảnh → Layout Options chọn “In line with text” để ảnh không "
       "trôi lung tung; căn giữa bằng Ctrl+E.")
bullet("Đặt chú thích tự động: chuột phải ảnh → Insert Caption → Label “Hình” để Word tự "
       "đánh số và tạo được Danh mục hình.")
bullet("Ảnh chụp màn hình điện thoại nên cắt gọn viền, độ phân giải cao; ảnh sơ đồ nên xuất "
       "dạng PNG nền trắng.")
bullet("Mỗi ảnh nên kèm một câu mô tả ngắn ngay dưới và được nhắc tới trong đoạn văn (ví dụ "
       "“… như minh họa ở Hình 6.1”).")

h2("Phụ lục B. Hướng dẫn cài đặt và sử dụng (tóm tắt)")
para("Yêu cầu thiết bị: Android 8.0 (API 26) trở lên, có camera sau và cảm biến gia tốc, "
     "khoảng 200 MB trống.")
numlist("Bật Tùy chọn nhà phát triển và Gỡ lỗi USB (hoặc cài trực tiếp tệp APK).")
numlist("Cài đặt qua lệnh gradlew installDebug hoặc chép tệp APK vào máy rồi mở để cài.")
numlist("Lần đầu mở: cấp quyền Camera và Micro; ứng dụng đọc lời chào hướng dẫn.")
numlist("Đếm tiền: đưa tờ tiền vào camera → nghe mệnh giá → vuốt xuống để cộng → chạm đôi "
        "nghe tổng → giữ lâu để đếm lại.")
numlist("Đọc menu: hướng camera vào menu → hệ thống tự chụp → nghe danh sách → vuốt phải/"
        "trái duyệt món → vuốt xuống chọn món.")
numlist("Dừng đọc bất kỳ lúc nào: lắc máy hoặc nhấn phím giảm âm lượng.")

h2("Phụ lục C. Cấu trúc mã nguồn")
code(
"app/src/main/kotlin/.../\n"
"  ui/        Các màn hình Compose (money, menu, history, settings)\n"
"  ml/        Nhận diện tiền (YOLOv10n), phân tích menu (VLM)\n"
"  tts/       Định tuyến TTS (FPT.AI + Android)\n"
"  voice/     Lệnh giọng nói, cảm biến lắc\n"
"  data/      Lịch sử (Room), cài đặt (DataStore)\n"
"  di/        Các module Hilt\n"
"  util/      Đọc số thành chữ tiếng Việt, rung\n"
"app/src/main/assets/ml/   Mô hình .tflite + nhãn\n"
"ml-training/notebooks/    Notebook huấn luyện trên Colab\n"
)
caption("Hình C.1. Cấu trúc thư mục mã nguồn dự án")

h2("Phụ lục D. Bảng đối chiếu mục tiêu và kết quả")
table(
    ["Mục tiêu (Mục 1.2)", "Kết quả đạt được"],
    [
        ["Mô hình nhận diện 9 mệnh giá VND thời gian thực",
         "Đã triển khai YOLOv10n on-device, có đếm cộng dồn"],
        ["Mô-đun đọc thực đơn (món + giá)",
         "Đã triển khai bằng mô hình ngôn ngữ – thị giác, trả về JSON"],
        ["Mô hình tương tác accessibility-first",
         "Đã thiết kế bộ cử chỉ + giọng nói + rung thống nhất"],
        ["Tích hợp TTS tiếng Việt tự nhiên",
         "Đã tích hợp FPT.AI + Android, có định tuyến dự phòng"],
        ["Đóng gói APK + tài liệu + quy trình huấn luyện",
         "Đã có APK chạy thật, tài liệu và notebook Colab tái lập được"],
    ],
    widths=[6.5, 8.0],
    caption_text="Bảng D.1. Đối chiếu mục tiêu đề tài và kết quả",
)

h2("Phụ lục E. Một số ca kiểm thử tiêu biểu")
table(
    ["Mã", "Mô tả", "Kỳ vọng", "Kết quả"],
    [
        ["TC-01", "Đưa tờ 200.000đ vào camera", "Đọc “Hai trăm nghìn đồng”", "Đạt"],
        ["TC-02", "Đếm 3 tờ 200k + 1 tờ 50k", "Tổng “Sáu trăm năm mươi nghìn đồng”", "Đạt"],
        ["TC-03", "Chụp menu 12 món", "Đọc “Menu có 12 món…”", "Đạt"],
        ["TC-04", "Mất mạng khi đọc menu", "Thông báo lỗi mạng, cho thử lại", "Đạt"],
        ["TC-05", "Lắc máy khi đang đọc", "Dừng đọc ngay (panic stop)", "Đạt"],
        ["TC-06", "Lệnh giọng nói “đọc tiền”", "Mở chế độ đếm tiền", "Đạt"],
        ["TC-07", "Camera chĩa vào vật không phải tiền", "Không báo nhầm mệnh giá", "[…]"],
    ],
    widths=[1.6, 5.4, 4.5, 3.0],
    caption_text="Bảng E.1. Một số ca kiểm thử tiêu biểu",
)
para("(Các ô “[…]” cần điền/kiểm chứng thêm trong quá trình kiểm thử cuối.)", italic=True)

# ============================================================
#  X. TÀI LIỆU THAM KHẢO
# ============================================================
h1("X. Tài liệu và nguồn tham khảo")
refs = [
    "A. Wang, et al., “YOLOv10: Real-Time End-to-End Object Detection,” arXiv:2405.14458, 2024.",
    "A. Howard, et al., “Searching for MobileNetV3,” Proc. IEEE/CVF ICCV, 2019.",
    "Meta AI, “The Llama 4 herd: native multimodal models,” 2025. "
    "https://ai.meta.com/blog/llama-4/",
    "Google, “TensorFlow Lite – On-device machine learning,” https://www.tensorflow.org/lite",
    "Google, “ML Kit Text Recognition v2,” "
    "https://developers.google.com/ml-kit/vision/text-recognition/v2",
    "Google, “Android Developers – Jetpack Compose, CameraX, Room, Hilt,” "
    "https://developer.android.com",
    "Groq, “GroqCloud Documentation,” https://console.groq.com/docs",
    "FPT.AI, “Text to Speech API,” https://fpt.ai/tts",
    "Ultralytics, “YOLO Documentation,” https://docs.ultralytics.com",
    "World Health Organization, “Blindness and vision impairment,” 2023. https://www.who.int",
    "J. Brooke, “SUS: A quick and dirty usability scale,” Usability Evaluation in Industry, 1996.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    p.paragraph_format.space_after = Pt(6)
    _font(p.add_run(f"[{i}] {r}"), 12)

root = ROOT
out = os.path.join(root, "BaoCao_DoAn_MatAI_co_hinh.docx")
try:
    doc.save(out)
except PermissionError:
    out = os.path.join(root, "BaoCao_DoAn_MatAI_co_hinh_v2.docx")
    doc.save(out)
    print("(File dang mo trong Word — luu ban moi)")

# Thong ke uoc luong so trang
words = 0
for p in doc.paragraphs:
    words += len(p.text.split())
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            words += len(c.text.split())
print("Saved:", out)
print("Tong so doan:", len(doc.paragraphs), "| Bang:", len(doc.tables), "| ~Tu:", words)
print("Uoc luong so trang (~270 tu/trang):", round(words / 270))
